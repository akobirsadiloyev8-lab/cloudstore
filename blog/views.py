from django.shortcuts import render, get_object_or_404
from django.db import models

# Birinchi all_books va adabiyotlar 700-qatorda mavjud

def book_list(request, author_id):
    from .models import Author, Book
    
    author = get_object_or_404(Author, id=author_id)
    books = Book.objects.filter(author=author)
    return render(request, 'blog/book_list.html', {'author': author, 'books': books})


from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import os

# Groq AI sozlash (xavfsiz import)
GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
groq_client = None
try:
    from groq import Groq
    if GROQ_API_KEY:
        groq_client = Groq(api_key=GROQ_API_KEY)
except Exception:
    # Import yoki versiya xatosi bo'lsa o'tkazib yuborish
    groq_client = None

@csrf_exempt
def ai_search_books(request):
    """Kitoblar ichidan AI bilan aqlli qidiruv"""
    from .models import Book, Author, BookPage
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST so\'rov qabul qilinadi'}, status=400)
    
    try:
        data = json.loads(request.body)
        original_query = data.get('query', '').strip()
        author_id = data.get('author_id')
        
        if not original_query:
            return JsonResponse({'error': 'Qidiruv so\'rovi bo\'sh'}, status=400)
        
        # Kitoblarni olish
        if author_id:
            books = Book.objects.filter(author_id=author_id)
        else:
            books = Book.objects.all()
        
        if not books.exists():
            return JsonResponse({
                'success': False,
                'message': 'Kitoblar topilmadi'
            })
        
        # 1-BOSQICH: AI so'rovni tahlil qilsin
        search_terms = [original_query.lower()]  # Asosiy so'rov
        ai_analysis = None
        
        if groq_client:
            try:
                analysis_prompt = f"""Foydalanuvchi qidiruv so'rovi: "{original_query}"

Vazifang:
1. Imloviy xatolarni tuzat (masalan: "muhabbat" -> "muhabbat", "kitop" -> "kitob")
2. Agar bu savol bo'lsa - qidirilishi kerak bo'lgan kalit so'zlarni ajrat
3. Agar bu gap bo'lsa - asosiy iboralarni ajrat
4. Sinonimlar yoki yaqin ma'noli so'zlarni qo'sh

MUHIM: Faqat JSON formatda javob ber, boshqa hech narsa yozma:
{{"corrected": "tuzatilgan so'rov", "keywords": ["kalit1", "kalit2", "kalit3"], "is_question": true/false}}"""

                analysis_completion = groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "Sen matn tahlilchisisisan. Faqat JSON formatda javob ber."},
                        {"role": "user", "content": analysis_prompt}
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=200
                )
                
                ai_result = analysis_completion.choices[0].message.content.strip()
                # JSON ni parse qilish
                import re
                json_match = re.search(r'\{[^{}]*\}', ai_result)
                if json_match:
                    ai_analysis = json.loads(json_match.group())
                    # Tuzatilgan so'rovni qo'shish
                    if ai_analysis.get('corrected'):
                        corrected = ai_analysis['corrected'].lower()
                        if corrected != original_query.lower() and corrected not in search_terms:
                            search_terms.insert(0, corrected)
                    # Kalit so'zlarni qo'shish
                    if ai_analysis.get('keywords'):
                        for kw in ai_analysis['keywords']:
                            kw_lower = kw.lower().strip()
                            if kw_lower and kw_lower not in search_terms and len(kw_lower) > 2:
                                search_terms.append(kw_lower)
            except Exception as e:
                pass
        
        # 2-BOSQICH: Barcha kalit so'zlar bo'yicha qidirish
        all_found_snippets = []
        found_positions = set()  # Takrorlanishni oldini olish
        relevant_texts = []
        
        for search_term in search_terms[:5]:  # Maksimum 5 ta term
            for book in books:
                pages = book.pages.all()
                for page in pages:
                    text_lower = (page.text or '').lower()
                    start_pos = 0
                    while True:
                        idx = text_lower.find(search_term, start_pos)
                        if idx == -1:
                            break
                        
                        # Takrorlanishni tekshirish
                        position_key = f"{book.id}-{page.page_number}-{idx}"
                        if position_key in found_positions:
                            start_pos = idx + 1
                            continue
                        found_positions.add(position_key)
                        
                        snippet_start = max(0, idx - 150)
                        snippet_end = min(len(page.text), idx + len(search_term) + 150)
                        snippet = page.text[snippet_start:snippet_end]
                        highlight_start = idx - snippet_start
                        highlight_end = highlight_start + len(search_term)
                        highlighted_snippet = (
                            snippet[:highlight_start] +
                            f'<mark>{snippet[highlight_start:highlight_end]}</mark>' +
                            snippet[highlight_end:]
                        )
                        all_found_snippets.append({
                            'book_id': book.id,
                            'book_title': book.title,
                            'author_name': book.author.name,
                            'snippet': f"...{highlighted_snippet}...",
                            'page_number': page.page_number,
                            'position': idx,
                            'search_term': search_term
                        })
                        if len(relevant_texts) < 5:
                            relevant_texts.append(f"[{book.title}, {page.page_number}-sahifa]: {snippet}")
                        start_pos = idx + 1
        
        # 3-BOSQICH: Agar savol bo'lsa, AI javob bersin
        ai_response = None
        is_question = (ai_analysis and ai_analysis.get('is_question')) or original_query.endswith('?') or original_query.lower().startswith(('nima', 'kim', 'qanday', 'qachon', 'nega', 'qaysi', 'qayer', 'necha', 'qancha', 'nimaga'))
        
        if groq_client and is_question and relevant_texts:
            try:
                context = "\n\n".join(relevant_texts)
                answer_prompt = f"""Savol: "{original_query}"

Kitoblardan topilgan ma'lumotlar:
{context}

Vazifa: Shu ma'lumotlar asosida savolga qisqa javob ber (2-3 gap). O'zbek tilida."""

                answer_completion = groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "Sen yordamchi assistantsan. Savolga qisqa javob ber."},
                        {"role": "user", "content": answer_prompt}
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.3,
                    max_tokens=300
                )
                ai_response = answer_completion.choices[0].message.content
            except:
                pass
        
        # Natijalarni tartiblash
        all_found_snippets.sort(key=lambda x: (x['book_title'], x['page_number'], x['position']))
        total_occurrences = len(all_found_snippets)
        unique_books = len(set(s['book_id'] for s in all_found_snippets))
        
        # Qidirilgan so'zlar haqida ma'lumot
        search_info = None
        if ai_analysis and ai_analysis.get('corrected') and ai_analysis['corrected'].lower() != original_query.lower():
            search_info = f"Tuzatildi: \"{ai_analysis['corrected']}\""
        if len(search_terms) > 1:
            search_info = (search_info + " | " if search_info else "") + f"Qidirildi: {', '.join(search_terms[:3])}"
        
        return JsonResponse({
            'success': True,
            'results': all_found_snippets[:20],
            'ai_response': ai_response,
            'search_info': search_info,
            'message': f"'{original_query}' bo'yicha {total_occurrences} ta natija topildi ({unique_books} ta kitobda)" if total_occurrences > 0 else f"'{original_query}' kitoblardan topilmadi",
            'total_occurrences': total_occurrences,
            'unique_books': unique_books
        })
            
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Noto\'g\'ri JSON format'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
import os
import sys
import tempfile
import traceback
import json
import re
import time
import shutil
import subprocess
import base64
import html as html_module
from django.shortcuts import render, redirect
from django.http import HttpResponse, Http404, JsonResponse
from django.core.files.storage import FileSystemStorage
from django.contrib.auth import login, authenticate
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.decorators import login_required
from .forms import ImageForm
from .models import Image, Feedback, File

# LibreOffice konverter - Word, Excel, PowerPoint uchun
try:
    # Asosiy papkadan import qilish
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from libreoffice_converter import (
        convert_to_pdf_with_libreoffice,
        convert_libreoffice_format,
        is_libreoffice_available
    )
except ImportError as e:
    print(f"LibreOffice converter import xatosi: {e}")
    # Agar import qilib bo'lmasa, None qilib qo'yamiz
    convert_to_pdf_with_libreoffice = None
    convert_libreoffice_format = None
    is_libreoffice_available = lambda: False

# Universal PDF converter (zaxira usul)
try:
    # Root papkadan import qilish
    import sys
    from pathlib import Path
    root_path = str(Path(__file__).resolve().parent.parent.parent)
    if root_path not in sys.path:
        sys.path.insert(0, root_path)
    from universal_pdf_converter import convert_any_to_pdf
except Exception as e:
    print(f"universal_pdf_converter import xatosi: {e}")
    convert_any_to_pdf = None

# PDF o'qish uchun advanced function
def extract_pdf_text_advanced(filepath):
    """PDF matnini turli usullar bilan o'qish (fitz, pdfplumber, PyPDF2, OCR)"""
    text = ""
    
    # 1-usul: PyMuPDF (fitz) - eng tez
    try:
        import fitz
        doc = fitz.open(filepath)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        if text and len(text.strip()) >= 50:
            return text.strip()
    except Exception as e:
        print(f"PyMuPDF xatosi: {e}")
    
    # 2-usul: pdfplumber - ko'proq matn topadi
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        if text and len(text.strip()) >= 50:
            return text.strip()
    except Exception as e:
        print(f"pdfplumber xatosi: {e}")
    
    # 3-usul: PyPDF2 - boshqa format
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text = "\n".join([page.extract_text() for page in reader.pages])
        if text and len(text.strip()) >= 50:
            return text.strip()
    except Exception as e:
        print(f"PyPDF2 xatosi: {e}")
    
    # 4-usul: OCR (rasm matnlarni o'qish)
    try:
        import pytesseract
        from pdf2image import convert_from_path
        from PIL import Image
        
        # PDF'ni rasmlarga aylantirish
        pages = convert_from_path(filepath, dpi=200, first_page=1, last_page=5)
        ocr_text = ""
        
        for i, page in enumerate(pages):
            page_text = pytesseract.image_to_string(page, lang='eng+uzb+rus')
            ocr_text += f"\n--- Sahifa {i+1} ---\n{page_text}\n"
        
        if ocr_text and len(ocr_text.strip()) >= 20:
            return ocr_text.strip()
            
    except ImportError:
        return "OCR kutubxonalari o'rnatilmagan (pytesseract, pdf2image kerak)"
    except Exception as e:
        return f"OCR xatosi: {str(e)}"
    
    return 'PDF faylidan matn oqib bolmadi. Fayl shikastlangan yoki matn yoq bolishi mumkin.'


# agar docx/pdf ishlatmoqchi bo'lsangiz, quyidagilarni import qiling:
import docx
from pypdf import PdfReader
from io import BytesIO
import zipfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader

from odf import text as odf_text, teletype
from odf.opendocument import OpenDocumentText
from ebooklib import epub

import openpyxl
import xlwt
from PIL import Image as PILImage

# Groq va CloudConvert - xavfsiz import
try:
    from groq import Groq
    GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
    groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None
except Exception:
    # Import yoki versiya xatosi bo'lsa o'tkazib yuborish
    groq_client = None

try:
    import cloudconvert
    CLOUDCONVERT_API_KEY = os.environ.get('CLOUDCONVERT_API_KEY', '')
    if CLOUDCONVERT_API_KEY:
        cloudconvert.configure(api_key=CLOUDCONVERT_API_KEY)
except Exception:
    cloudconvert = None
    CLOUDCONVERT_API_KEY = ''


def boshlash(request):
    context = {}
    if request.method == "POST":
        if 'upload_image' in request.POST:
            form = ImageForm(request.POST, request.FILES)
            if form.is_valid():
                form.instance.user = request.user
                form.save()
                return redirect('image_list')
        elif request.FILES.get("uploaded_file"):
            uploaded_file = request.FILES["uploaded_file"]

            file_instance = File(user=request.user, file=uploaded_file)
            file_instance.save()
            filename = file_instance.file.name
            return redirect('preview', filename=filename)

    return render(request, "blog/boshlash.html", context)


def preview_files(request):
    # Bazadan fayllarni olish
    from django.conf import settings
    
    existing_files = []
    
    # Bazadagi fayllar - faqat o'z foydalanuvchining fayllari
    if request.user.is_authenticated:
        db_files = File.objects.filter(user=request.user).order_by('-uploaded_at')
    else:
        # Login qilmagan foydalanuvchilar uchun faqat user=None bo'lgan fayllar
        db_files = File.objects.filter(user=None).order_by('-uploaded_at')
    
    for f in db_files:
        if f.file and os.path.exists(f.file.path):
            existing_files.append({
                'id': f.id,
                'name': os.path.basename(f.file.name),
                'path': f.file.path,
                'size': os.path.getsize(f.file.path),
                'uploaded_at': f.uploaded_at
            })
    
    return render(request, 'blog/malumot.html', {'existing_files': existing_files})


from django.views.decorators.csrf import csrf_exempt
import json
@csrf_exempt
def preview_file_content(request):
    # converter/views.py dagi preview_file_content funksiyasi
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            filepath = data.get('filepath')
            if not filepath or not os.path.exists(filepath):
                return JsonResponse({'success': False, 'error': 'Fayl topilmadi'})
            file_path = filepath
            ext = os.path.splitext(file_path)[1].lower()
            filename = os.path.basename(file_path)
            text_extensions = ['.txt', '.html', '.htm', '.css', '.js', '.py', '.json', '.xml', '.csv', '.md']
            image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']
            pdf_extensions = ['.pdf']
            result = {
                'success': True,
                'filename': filename,
                'file_type': 'unknown',
                'content': None,
                'preview_url': None
            }
            if ext in text_extensions:
                result['file_type'] = 'text'
                encodings = ['utf-8', 'cp1251', 'latin-1', 'utf-16']
                content = None
                for encoding in encodings:
                    try:
                        with open(filepath, 'r', encoding=encoding) as f:
                            content = f.read()
                        break
                    except:
                        continue
                if content is None:
                    with open(filepath, 'rb') as f:
                        content = f.read().decode('utf-8', errors='replace')
                result['content'] = content
            elif ext in image_extensions:
                result['file_type'] = 'image'
                import base64
                with open(filepath, 'rb') as f:
                    image_data = base64.b64encode(f.read()).decode('utf-8')
                mime_types = {
                    '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                    '.png': 'image/png', '.gif': 'image/gif',
                    '.bmp': 'image/bmp', '.webp': 'image/webp',
                    '.svg': 'image/svg+xml'
                }
                mime = mime_types.get(ext, 'image/png')
                result['content'] = f'data:{mime};base64,{image_data}'
            elif ext in pdf_extensions:
                result['file_type'] = 'pdf'
                # Faqat PyMuPDF (fitz) orqali PDF matnini o'qish
                try:
                    import fitz
                    doc = fitz.open(filepath)
                    text = "\n".join([page.get_text() for page in doc])
                    if text and text.strip():
                        result['content'] = text
                    else:
                        result['content'] = 'PDF faylidan matn o‘qib bo‘lmadi yoki fayl bo‘sh.'
                except Exception as e:
                    result['content'] = f'PDF matnini o‘qishda xatolik: {str(e)}'
            elif ext in ['.doc', '.docx']:
                result['file_type'] = 'document'
                content = None
                try:
                    from docx import Document
                    from docx.shared import Pt
                    doc = Document(filepath)
                    
                    # Word faylini HTML formatiga o'tkazish
                    html_parts = ['<div class="word-document">']
                    
                    for element in doc.element.body:
                        if element.tag.endswith('p'):  # Paragraf
                            for para in doc.paragraphs:
                                if para._element == element:
                                    style_name = para.style.name if para.style else ''
                                    text = para.text.strip()
                                    if not text:
                                        continue
                                    
                                    # Sarlavhalarni aniqlash
                                    if 'Heading 1' in style_name or style_name == 'Title':
                                        html_parts.append(f'<h1>{text}</h1>')
                                    elif 'Heading 2' in style_name:
                                        html_parts.append(f'<h2>{text}</h2>')
                                    elif 'Heading 3' in style_name:
                                        html_parts.append(f'<h3>{text}</h3>')
                                    elif 'Heading' in style_name:
                                        html_parts.append(f'<h4>{text}</h4>')
                                    else:
                                        # Oddiy paragraf - bold/italic tekshirish
                                        formatted_text = ''
                                        for run in para.runs:
                                            run_text = run.text
                                            if run.bold and run.italic:
                                                formatted_text += f'<strong><em>{run_text}</em></strong>'
                                            elif run.bold:
                                                formatted_text += f'<strong>{run_text}</strong>'
                                            elif run.italic:
                                                formatted_text += f'<em>{run_text}</em>'
                                            else:
                                                formatted_text += run_text
                                        if formatted_text.strip():
                                            html_parts.append(f'<p>{formatted_text}</p>')
                                    break
                        
                        elif element.tag.endswith('tbl'):  # Jadval
                            for table in doc.tables:
                                if table._element == element:
                                    html_parts.append('<table class="word-table">')
                                    for i, row in enumerate(table.rows):
                                        html_parts.append('<tr>')
                                        for cell in row.cells:
                                            tag = 'th' if i == 0 else 'td'
                                            html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                                        html_parts.append('</tr>')
                                    html_parts.append('</table>')
                                    break
                    
                    html_parts.append('</div>')
                    content = '\n'.join(html_parts)
                    result['file_type'] = 'html_document'
                    
                except Exception as e:
                    # Oddiy matn sifatida o'qish
                    try:
                        from docx import Document
                        doc = Document(filepath)
                        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                        content = '\n\n'.join(paragraphs)
                    except:
                        pass
                
                if not content:
                    try:
                        import textract
                        content = textract.process(str(filepath)).decode('utf-8', errors='replace')
                    except Exception as e:
                        content = "Word faylini o'qib bo'lmadi. python-docx yoki textract kutubxonasi kerak."
                result['content'] = content
            elif ext in ['.xls', '.xlsx']:
                result['file_type'] = 'spreadsheet'
                try:
                    import pandas as pd
                    df = pd.read_excel(filepath)
                    result['content'] = df.to_html(classes='excel-table', index=False)
                except:
                    result['content'] = 'Excel faylini o\'qib bo\'lmadi. pandas va openpyxl kerak.'
            else:
                result['file_type'] = 'binary'
                result['content'] = 'Bu fayl turini ko\'rib bo\'lmaydi.'
            return JsonResponse(result)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'POST kerak'})


# Upload va convert funksiyalari
from pathlib import Path
from django.conf import settings

TEMP_DIR = Path(settings.BASE_DIR) / 'temp_files'
TEMP_DIR.mkdir(exist_ok=True)
OUTPUT_DIR = Path(settings.BASE_DIR) / 'converted_files'
OUTPUT_DIR.mkdir(exist_ok=True)


@csrf_exempt
def upload_file(request):
    if request.method == 'POST':
        try:
            uploaded_file = request.FILES.get('file')
            if uploaded_file:
                # Fayl saqlash logikasi
                file_path = os.path.join(settings.MEDIA_ROOT, 'files', uploaded_file.name)
                with open(file_path, 'wb+') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)
                return JsonResponse({'success': True, 'message': 'Fayl yuklandi'})
            else:
                return JsonResponse({'success': False, 'error': 'Fayl tanlanmadi'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@csrf_exempt
def convert_to_format(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            filepath = data.get('filepath')
            output_format = data.get('format', 'pdf')
            
            if not filepath or not os.path.exists(filepath):
                return JsonResponse({'success': False, 'error': 'Fayl topilmadi'})
            
            input_path = Path(filepath)
            result = None
            ext = input_path.suffix.lower()
            output_path = OUTPUT_DIR / f"{input_path.stem}.{output_format}"
            
            # PDF ga konvertatsiya
            if output_format == 'pdf':
                output_path = OUTPUT_DIR / f"{input_path.stem}.pdf"
                
                # Python kutubxonalari orqali konvertatsiya (LibreOffice o'rniga)
                if ext == '.docx':
                    try:
                        # docx faylni matn sifatida o'qish va PDF yaratish
                        from docx import Document
                        from reportlab.pdfgen import canvas
                        from reportlab.lib.pagesizes import A4
                        
                        # DOCX'dan matn o'qish
                        doc = Document(input_path)
                        text = "\n".join([p.text for p in doc.paragraphs])
                        
                        # PDF yaratish
                        c = canvas.Canvas(str(output_path), pagesize=A4)
                        width, height = A4
                        
                        # Matnni PDF'ga yozish
                        y_position = height - 50
                        for line in text.split('\n'):
                            if y_position < 50:
                                c.showPage()
                                y_position = height - 50
                            if line.strip():  # Bo'sh qatorlarni o'tkazib yuborish
                                c.drawString(50, y_position, line[:100])  # 100 belgigacha
                                y_position -= 20
                        
                        c.save()
                        result = str(output_path)
                        print(f"✓ DOCX dan PDF yaratildi: {result}")
                    except Exception as e:
                        result = f"DOCX konvertatsiya xatosi: {str(e)}"
                        print(f"✗ DOCX konvertatsiya xatosi: {str(e)}")
                else:
                    result = f"LibreOffice o'rnatilmagan. Faqat DOCX -> PDF qo'llab-quvvatlanadi."
                
                # Natijani qaytarish
                if result:
                    return JsonResponse({'success': True, 'output_file': result})
                else:
                    return JsonResponse({'success': False, 'error': 'Konvertatsiya amalga oshmadi'})
                    
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


def adabiyotlar(request):
    from .models import Author
    authors = Author.objects.all()
    return render(request, 'blog/adabiyotlar.html', {'authors': authors})

def all_books(request):
    """Barcha kitoblarni ko'rsatish - tab tizimi bilan birlashtirilgan"""
    from .models import Book, Category, Author, Favorite, SearchQuery
    from django.db.models import Avg
    
    categories = Category.objects.all()
    authors = Author.objects.all()
    
    # Tab tizimi
    current_tab = request.GET.get('tab', 'all')
    
    if current_tab == 'top':
        # Eng ko'p ko'rilgan va yuqori baholangan
        books = Book.objects.annotate(
            avg_rating=Avg('ratings__rating')
        ).order_by('-views_count', '-avg_rating')
    elif current_tab == 'new':
        # Yangi qo'shilgan kitoblar
        books = Book.objects.order_by('-created_at')
    elif current_tab == 'favorites':
        # Foydalanuvchi sevimlilari
        if request.user.is_authenticated:
            favorite_ids = Favorite.objects.filter(user=request.user).values_list('book_id', flat=True)
            books = Book.objects.filter(id__in=favorite_ids)
        else:
            books = Book.objects.none()
    else:
        # Barcha kitoblar
        books = Book.objects.all().order_by('-created_at')
    
    books = books.select_related('author').prefetch_related('categories')
    
    # Filter by category
    category_slug = request.GET.get('category')
    if category_slug:
        books = books.filter(categories__slug=category_slug)
    
    # Filter by author
    author_id = request.GET.get('author')
    if author_id:
        books = books.filter(author_id=author_id)
    
    # Search
    search = request.GET.get('q')
    if search:
        books = books.filter(
            models.Q(title__icontains=search) | 
            models.Q(author__name__icontains=search)
        )
    
    # Top searches for top tab
    top_searches = []
    if current_tab == 'top':
        top_searches = SearchQuery.objects.all()[:10]
    
    return render(request, 'blog/all_books.html', {
        'books': books,
        'categories': categories,
        'authors': authors,
        'current_category': category_slug,
        'current_author': int(author_id) if author_id else None,
        'search_query': search,
        'current_tab': current_tab,
        'top_searches': top_searches,
    })

# top_books, new_books, favorites_list - all_books ga birlashtirildi (redirect sifatida)
def top_books(request):
    """all_books ga yo'naltirish"""
    from django.shortcuts import redirect
    return redirect('/kitoblar/?tab=top')

def new_books(request):
    """all_books ga yo'naltirish"""
    from django.shortcuts import redirect
    return redirect('/kitoblar/?tab=new')

def favorites_list(request):
    """all_books ga yo'naltirish"""
    from django.shortcuts import redirect
    return redirect('/kitoblar/?tab=favorites')

def book_list(request, author_id):
    from .models import Author, Book
    
    author = get_object_or_404(Author, id=author_id)
    books = Book.objects.filter(author=author)
    return render(request, 'blog/book_list.html', {'author': author, 'books': books})


from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import os

# Groq AI sozlash (xavfsiz import)
GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
groq_client = None
try:
    from groq import Groq
    if GROQ_API_KEY:
        groq_client = Groq(api_key=GROQ_API_KEY)
except Exception:
    # Import yoki versiya xatosi bo'lsa o'tkazib yuborish
    groq_client = None

@csrf_exempt
def ai_search_books(request):
    """Kitoblar ichidan AI bilan aqlli qidiruv"""
    from .models import Book, Author, BookPage
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST so\'rov qabul qilinadi'}, status=400)
    
    try:
        data = json.loads(request.body)
        original_query = data.get('query', '').strip()
        author_id = data.get('author_id')
        
        if not original_query:
            return JsonResponse({'error': 'Qidiruv so\'rovi bo\'sh'}, status=400)
        
        # Kitoblarni olish
        if author_id:
            books = Book.objects.filter(author_id=author_id)
        else:
            books = Book.objects.all()
        
        if not books.exists():
            return JsonResponse({
                'success': False,
                'message': 'Kitoblar topilmadi'
            })
        
        # 1-BOSQICH: AI so'rovni tahlil qilsin
        search_terms = [original_query.lower()]  # Asosiy so'rov
        ai_analysis = None
        
        if groq_client:
            try:
                analysis_prompt = f"""Foydalanuvchi qidiruv so'rovi: "{original_query}"

Vazifang:
1. Imloviy xatolarni tuzat (masalan: "muhabbat" -> "muhabbat", "kitop" -> "kitob")
2. Agar bu savol bo'lsa - qidirilishi kerak bo'lgan kalit so'zlarni ajrat
3. Agar bu gap bo'lsa - asosiy iboralarni ajrat
4. Sinonimlar yoki yaqin ma'noli so'zlarni qo'sh

MUHIM: Faqat JSON formatda javob ber, boshqa hech narsa yozma:
{{"corrected": "tuzatilgan so'rov", "keywords": ["kalit1", "kalit2", "kalit3"], "is_question": true/false}}"""

                analysis_completion = groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "Sen matn tahlilchisisisan. Faqat JSON formatda javob ber."},
                        {"role": "user", "content": analysis_prompt}
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.1,
                    max_tokens=200
                )
                
                ai_result = analysis_completion.choices[0].message.content.strip()
                # JSON ni parse qilish
                import re
                json_match = re.search(r'\{[^{}]*\}', ai_result)
                if json_match:
                    ai_analysis = json.loads(json_match.group())
                    # Tuzatilgan so'rovni qo'shish
                    if ai_analysis.get('corrected'):
                        corrected = ai_analysis['corrected'].lower()
                        if corrected != original_query.lower() and corrected not in search_terms:
                            search_terms.insert(0, corrected)
                    # Kalit so'zlarni qo'shish
                    if ai_analysis.get('keywords'):
                        for kw in ai_analysis['keywords']:
                            kw_lower = kw.lower().strip()
                            if kw_lower and kw_lower not in search_terms and len(kw_lower) > 2:
                                search_terms.append(kw_lower)
            except Exception as e:
                pass
        
        # 2-BOSQICH: Barcha kalit so'zlar bo'yicha qidirish
        all_found_snippets = []
        found_positions = set()  # Takrorlanishni oldini olish
        relevant_texts = []
        
        for search_term in search_terms[:5]:  # Maksimum 5 ta term
            for book in books:
                pages = book.pages.all()
                for page in pages:
                    text_lower = (page.text or '').lower()
                    start_pos = 0
                    while True:
                        idx = text_lower.find(search_term, start_pos)
                        if idx == -1:
                            break
                        
                        # Takrorlanishni tekshirish
                        position_key = f"{book.id}-{page.page_number}-{idx}"
                        if position_key in found_positions:
                            start_pos = idx + 1
                            continue
                        found_positions.add(position_key)
                        
                        snippet_start = max(0, idx - 150)
                        snippet_end = min(len(page.text), idx + len(search_term) + 150)
                        snippet = page.text[snippet_start:snippet_end]
                        highlight_start = idx - snippet_start
                        highlight_end = highlight_start + len(search_term)
                        highlighted_snippet = (
                            snippet[:highlight_start] +
                            f'<mark>{snippet[highlight_start:highlight_end]}</mark>' +
                            snippet[highlight_end:]
                        )
                        all_found_snippets.append({
                            'book_id': book.id,
                            'book_title': book.title,
                            'author_name': book.author.name,
                            'snippet': f"...{highlighted_snippet}...",
                            'page_number': page.page_number,
                            'position': idx,
                            'search_term': search_term
                        })
                        if len(relevant_texts) < 5:
                            relevant_texts.append(f"[{book.title}, {page.page_number}-sahifa]: {snippet}")
                        start_pos = idx + 1
        
        # 3-BOSQICH: Agar savol bo'lsa, AI javob bersin
        ai_response = None
        is_question = (ai_analysis and ai_analysis.get('is_question')) or original_query.endswith('?') or original_query.lower().startswith(('nima', 'kim', 'qanday', 'qachon', 'nega', 'qaysi', 'qayer', 'necha', 'qancha', 'nimaga'))
        
        if groq_client and is_question and relevant_texts:
            try:
                context = "\n\n".join(relevant_texts)
                answer_prompt = f"""Savol: "{original_query}"

Kitoblardan topilgan ma'lumotlar:
{context}

Vazifa: Shu ma'lumotlar asosida savolga qisqa javob ber (2-3 gap). O'zbek tilida."""

                answer_completion = groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "Sen yordamchi assistantsan. Savolga qisqa javob ber."},
                        {"role": "user", "content": answer_prompt}
                    ],
                    model="llama-3.1-8b-instant",
                    temperature=0.3,
                    max_tokens=300
                )
                ai_response = answer_completion.choices[0].message.content
            except:
                pass
        
        # Natijalarni tartiblash
        all_found_snippets.sort(key=lambda x: (x['book_title'], x['page_number'], x['position']))
        total_occurrences = len(all_found_snippets)
        unique_books = len(set(s['book_id'] for s in all_found_snippets))
        
        # Qidirilgan so'zlar haqida ma'lumot
        search_info = None
        if ai_analysis and ai_analysis.get('corrected') and ai_analysis['corrected'].lower() != original_query.lower():
            search_info = f"Tuzatildi: \"{ai_analysis['corrected']}\""
        if len(search_terms) > 1:
            search_info = (search_info + " | " if search_info else "") + f"Qidirildi: {', '.join(search_terms[:3])}"
        
        return JsonResponse({
            'success': True,
            'results': all_found_snippets[:20],
            'ai_response': ai_response,
            'search_info': search_info,
            'message': f"'{original_query}' bo'yicha {total_occurrences} ta natija topildi ({unique_books} ta kitobda)" if total_occurrences > 0 else f"'{original_query}' kitoblardan topilmadi",
            'total_occurrences': total_occurrences,
            'unique_books': unique_books
        })
            
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Noto\'g\'ri JSON format'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
import os
import sys
import tempfile
import traceback
import json
import re
import time
import shutil
import subprocess
import base64
import html as html_module
from django.shortcuts import render, redirect
from django.http import HttpResponse, Http404, JsonResponse
from django.core.files.storage import FileSystemStorage
from django.contrib.auth import login, authenticate
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.decorators import login_required
from .forms import ImageForm
from .models import Image, Feedback, File

# LibreOffice konverter - Word, Excel, PowerPoint uchun
try:
    # Asosiy papkadan import qilish
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from libreoffice_converter import (
        convert_to_pdf_with_libreoffice,
        convert_libreoffice_format,
        is_libreoffice_available
    )
except ImportError as e:
    print(f"LibreOffice converter import xatosi: {e}")
    # Agar import qilib bo'lmasa, None qilib qo'yamiz
    convert_to_pdf_with_libreoffice = None
    convert_libreoffice_format = None
    is_libreoffice_available = lambda: False

# Universal PDF converter (zaxira usul)
try:
    # Root papkadan import qilish
    import sys
    from pathlib import Path
    root_path = str(Path(__file__).resolve().parent.parent.parent)
    if root_path not in sys.path:
        sys.path.insert(0, root_path)
    from universal_pdf_converter import convert_any_to_pdf
except Exception as e:
    print(f"universal_pdf_converter import xatosi: {e}")
    convert_any_to_pdf = None

# PDF o'qish uchun advanced function
def extract_pdf_text_advanced(filepath):
    """PDF matnini turli usullar bilan o'qish (fitz, pdfplumber, PyPDF2, OCR)"""
    text = ""
    
    # 1-usul: PyMuPDF (fitz) - eng tez
    try:
        import fitz
        doc = fitz.open(filepath)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        if text and len(text.strip()) >= 50:
            return text.strip()
    except Exception as e:
        print(f"PyMuPDF xatosi: {e}")
    
    # 2-usul: pdfplumber - ko'proq matn topadi
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        if text and len(text.strip()) >= 50:
            return text.strip()
    except Exception as e:
        print(f"pdfplumber xatosi: {e}")
    
    # 3-usul: PyPDF2 - boshqa format
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text = "\n".join([page.extract_text() for page in reader.pages])
        if text and len(text.strip()) >= 50:
            return text.strip()
    except Exception as e:
        print(f"PyPDF2 xatosi: {e}")
    
    # 4-usul: OCR (rasm matnlarni o'qish)
    try:
        import pytesseract
        from pdf2image import convert_from_path
        from PIL import Image
        
        # PDF'ni rasmlarga aylantirish
        pages = convert_from_path(filepath, dpi=200, first_page=1, last_page=5)
        ocr_text = ""
        
        for i, page in enumerate(pages):
            page_text = pytesseract.image_to_string(page, lang='eng+uzb+rus')
            ocr_text += f"\n--- Sahifa {i+1} ---\n{page_text}\n"
        
        if ocr_text and len(ocr_text.strip()) >= 20:
            return ocr_text.strip()
            
    except ImportError:
        return "OCR kutubxonalari o'rnatilmagan (pytesseract, pdf2image kerak)"
    except Exception as e:
        return f"OCR xatosi: {str(e)}"
    
    return 'PDF faylidan matn oqib bolmadi. Fayl shikastlangan yoki matn yoq bolishi mumkin.'


# agar docx/pdf ishlatmoqchi bo'lsangiz, quyidagilarni import qiling:
import docx
from pypdf import PdfReader
from io import BytesIO
import zipfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader

from odf import text as odf_text, teletype
from odf.opendocument import OpenDocumentText
from ebooklib import epub

import openpyxl
import xlwt
from PIL import Image as PILImage

# Groq va CloudConvert - xavfsiz import
try:
    from groq import Groq
    GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
    groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None
except Exception:
    # Import yoki versiya xatosi bo'lsa o'tkazib yuborish
    groq_client = None

try:
    import cloudconvert
    CLOUDCONVERT_API_KEY = os.environ.get('CLOUDCONVERT_API_KEY', '')
    if CLOUDCONVERT_API_KEY:
        cloudconvert.configure(api_key=CLOUDCONVERT_API_KEY)
except Exception:
    cloudconvert = None
    CLOUDCONVERT_API_KEY = ''


def boshlash(request):
    context = {}
    if request.method == "POST":
        if 'upload_image' in request.POST:
            form = ImageForm(request.POST, request.FILES)
            if form.is_valid():
                form.instance.user = request.user
                form.save()
                return redirect('image_list')
        elif request.FILES.get("uploaded_file"):
            uploaded_file = request.FILES["uploaded_file"]

            file_instance = File(user=request.user, file=uploaded_file)
            file_instance.save()
            filename = file_instance.file.name
            return redirect('preview', filename=filename)

    return render(request, "blog/boshlash.html", context)


def preview_files(request):
    # Bazadan fayllarni olish
    from django.conf import settings
    
    existing_files = []
    
    # Bazadagi fayllar - faqat o'z foydalanuvchining fayllari
    if request.user.is_authenticated:
        db_files = File.objects.filter(user=request.user).order_by('-uploaded_at')
    else:
        # Login qilmagan foydalanuvchilar uchun faqat user=None bo'lgan fayllar
        db_files = File.objects.filter(user=None).order_by('-uploaded_at')
    
    for f in db_files:
        if f.file and os.path.exists(f.file.path):
            existing_files.append({
                'id': f.id,
                'name': os.path.basename(f.file.name),
                'path': f.file.path,
                'size': os.path.getsize(f.file.path),
                'uploaded_at': f.uploaded_at
            })
    
    return render(request, 'blog/malumot.html', {'existing_files': existing_files})


from django.views.decorators.csrf import csrf_exempt
import json
@csrf_exempt
def preview_file_content(request):
    # converter/views.py dagi preview_file_content funksiyasi
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            filepath = data.get('filepath')
            if not filepath or not os.path.exists(filepath):
                return JsonResponse({'success': False, 'error': 'Fayl topilmadi'})
            file_path = filepath
            ext = os.path.splitext(file_path)[1].lower()
            filename = os.path.basename(file_path)
            text_extensions = ['.txt', '.html', '.htm', '.css', '.js', '.py', '.json', '.xml', '.csv', '.md']
            image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']
            pdf_extensions = ['.pdf']
            result = {
                'success': True,
                'filename': filename,
                'file_type': 'unknown',
                'content': None,
                'preview_url': None
            }
            if ext in text_extensions:
                result['file_type'] = 'text'
                encodings = ['utf-8', 'cp1251', 'latin-1', 'utf-16']
                content = None
                for encoding in encodings:
                    try:
                        with open(filepath, 'r', encoding=encoding) as f:
                            content = f.read()
                        break
                    except:
                        continue
                if content is None:
                    with open(filepath, 'rb') as f:
                        content = f.read().decode('utf-8', errors='replace')
                result['content'] = content
            elif ext in image_extensions:
                result['file_type'] = 'image'
                import base64
                with open(filepath, 'rb') as f:
                    image_data = base64.b64encode(f.read()).decode('utf-8')
                mime_types = {
                    '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                    '.png': 'image/png', '.gif': 'image/gif',
                    '.bmp': 'image/bmp', '.webp': 'image/webp',
                    '.svg': 'image/svg+xml'
                }
                mime = mime_types.get(ext, 'image/png')
                result['content'] = f'data:{mime};base64,{image_data}'
            elif ext in pdf_extensions:
                result['file_type'] = 'pdf'
                # Faqat PyMuPDF (fitz) orqali PDF matnini o'qish
                try:
                    import fitz
                    doc = fitz.open(filepath)
                    text = "\n".join([page.get_text() for page in doc])
                    if text and text.strip():
                        result['content'] = text
                    else:
                        result['content'] = 'PDF faylidan matn o‘qib bo‘lmadi yoki fayl bo‘sh.'
                except Exception as e:
                    result['content'] = f'PDF matnini o‘qishda xatolik: {str(e)}'
            elif ext in ['.doc', '.docx']:
                result['file_type'] = 'document'
                content = None
                try:
                    from docx import Document
                    from docx.shared import Pt
                    doc = Document(filepath)
                    
                    # Word faylini HTML formatiga o'tkazish
                    html_parts = ['<div class="word-document">']
                    
                    for element in doc.element.body:
                        if element.tag.endswith('p'):  # Paragraf
                            for para in doc.paragraphs:
                                if para._element == element:
                                    style_name = para.style.name if para.style else ''
                                    text = para.text.strip()
                                    if not text:
                                        continue
                                    
                                    # Sarlavhalarni aniqlash
                                    if 'Heading 1' in style_name or style_name == 'Title':
                                        html_parts.append(f'<h1>{text}</h1>')
                                    elif 'Heading 2' in style_name:
                                        html_parts.append(f'<h2>{text}</h2>')
                                    elif 'Heading 3' in style_name:
                                        html_parts.append(f'<h3>{text}</h3>')
                                    elif 'Heading' in style_name:
                                        html_parts.append(f'<h4>{text}</h4>')
                                    else:
                                        # Oddiy paragraf - bold/italic tekshirish
                                        formatted_text = ''
                                        for run in para.runs:
                                            run_text = run.text
                                            if run.bold and run.italic:
                                                formatted_text += f'<strong><em>{run_text}</em></strong>'
                                            elif run.bold:
                                                formatted_text += f'<strong>{run_text}</strong>'
                                            elif run.italic:
                                                formatted_text += f'<em>{run_text}</em>'
                                            else:
                                                formatted_text += run_text
                                        if formatted_text.strip():
                                            html_parts.append(f'<p>{formatted_text}</p>')
                                    break
                        
                        elif element.tag.endswith('tbl'):  # Jadval
                            for table in doc.tables:
                                if table._element == element:
                                    html_parts.append('<table class="word-table">')
                                    for i, row in enumerate(table.rows):
                                        html_parts.append('<tr>')
                                        for cell in row.cells:
                                            tag = 'th' if i == 0 else 'td'
                                            html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                                        html_parts.append('</tr>')
                                    html_parts.append('</table>')
                                    break
                    
                    html_parts.append('</div>')
                    content = '\n'.join(html_parts)
                    result['file_type'] = 'html_document'
                    
                except Exception as e:
                    # Oddiy matn sifatida o'qish
                    try:
                        from docx import Document
                        doc = Document(filepath)
                        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                        content = '\n\n'.join(paragraphs)
                    except:
                        pass
                
                if not content:
                    try:
                        import textract
                        content = textract.process(str(filepath)).decode('utf-8', errors='replace')
                    except Exception as e:
                        content = "Word faylini o'qib bo'lmadi. python-docx yoki textract kutubxonasi kerak."
                result['content'] = content
            elif ext in ['.xls', '.xlsx']:
                result['file_type'] = 'spreadsheet'
                try:
                    import pandas as pd
                    df = pd.read_excel(filepath)
                    result['content'] = df.to_html(classes='excel-table', index=False)
                except:
                    result['content'] = 'Excel faylini o\'qib bo\'lmadi. pandas va openpyxl kerak.'
            else:
                result['file_type'] = 'binary'
                result['content'] = 'Bu fayl turini ko\'rib bo\'lmaydi.'
            return JsonResponse(result)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'POST kerak'})


# Upload va convert funksiyalari
from pathlib import Path
from django.conf import settings

TEMP_DIR = Path(settings.BASE_DIR) / 'temp_files'
TEMP_DIR.mkdir(exist_ok=True)
OUTPUT_DIR = Path(settings.BASE_DIR) / 'converted_files'
OUTPUT_DIR.mkdir(exist_ok=True)


@csrf_exempt
def upload_file(request):
    if request.method == 'POST':
        try:
            uploaded_file = request.FILES.get('file')
            if uploaded_file:
                # Fayl saqlash logikasi
                file_path = os.path.join(settings.MEDIA_ROOT, 'files', uploaded_file.name)
                with open(file_path, 'wb+') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)
                return JsonResponse({'success': True, 'message': 'Fayl yuklandi'})
            else:
                return JsonResponse({'success': False, 'error': 'Fayl tanlanmadi'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@csrf_exempt
def convert_to_format(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            filepath = data.get('filepath')
            output_format = data.get('format', 'pdf')
            
            if not filepath or not os.path.exists(filepath):
                return JsonResponse({'success': False, 'error': 'Fayl topilmadi'})
            
            input_path = Path(filepath)
            result = None
            ext = input_path.suffix.lower()
            output_path = OUTPUT_DIR / f"{input_path.stem}.{output_format}"
            
            # PDF ga konvertatsiya
            if output_format == 'pdf':
                output_path = OUTPUT_DIR / f"{input_path.stem}.pdf"
                
                # LibreOffice orqali konvertatsiya (eng yaxshi usul - asl formatni saqlaydi)
                libre_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    r"C:\Program Files\LibreOffice\program\soffice.com",
                    "/usr/bin/libreoffice",
                    "/usr/bin/soffice"
                ]
                
                libre_path = None
                for path in libre_paths:
                    if os.path.exists(path):
                        libre_path = path
                        break
                
                if libre_path and ext in ['.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp', '.rtf', '.txt', '.html', '.htm']:
                    try:
                        import subprocess
                        # LibreOffice bilan PDF ga aylantirish
                        result_process = subprocess.run([
                            libre_path,
                            '--headless',
                            '--convert-to', 'pdf',
                            '--outdir', str(OUTPUT_DIR),
                            str(input_path)
                        ], capture_output=True, text=True, timeout=120)
                        
                        # Natija faylni tekshirish
                        expected_pdf = OUTPUT_DIR / f"{input_path.stem}.pdf"
                        if expected_pdf.exists():
                            result = str(expected_pdf)
                            print(f"вњ“ LibreOffice konvertatsiya muvaffaqiyatli: {result}")
                        else:
                            print(f"LibreOffice xatolik: {result_process.stderr}")
                    except subprocess.TimeoutExpired:
                        print("LibreOffice timeout")
                    except Exception as e:
                        print(f"LibreOffice xatolik: {e}")
                
                # Agar LibreOffice ishlamasa, docx2pdf bilan sinab ko'ramiz
                if not result and ext in ['.docx', '.doc']:
                    try:
                        from docx2pdf import convert
                        convert(str(input_path), str(output_path))
                        if output_path.exists():
                            result = str(output_path)
                    except Exception as e:
                        print(f"docx2pdf xatolik: {e}")
                
                # Rasmlar uchun
                if not result and ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
                    try:
                        from PIL import Image as PILImage
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfgen import canvas as pdf_canvas
                        
                        img = PILImage.open(str(input_path))
                        if img.mode == 'RGBA':
                            img = img.convert('RGB')
                        
                        img_width, img_height = img.size
                        page_width, page_height = letter
                        
                        ratio = min(page_width / img_width, page_height / img_height) * 0.9
                        new_width = img_width * ratio
                        new_height = img_height * ratio
                        
                        c = pdf_canvas.Canvas(str(output_path), pagesize=letter)
                        x = (page_width - new_width) / 2
                        y = (page_height - new_height) / 2
                        c.drawImage(str(input_path), x, y, width=new_width, height=new_height)
                        c.save()
                        result = str(output_path)
                    except Exception as e:
                        print(f"Rasm->PDF xatolik: {e}")
                
                # Zaxira usul - convert_any_to_pdf
                if not result and convert_any_to_pdf:
                    try:
                        result = convert_any_to_pdf(str(input_path), str(output_path))
                    except Exception as e:
                        print(f"convert_any_to_pdf xatolik: {e}")
                
                # Eng oxirgi zaxira - matndan PDF
                if not result:
                    text_content = extract_text_content(input_path)
                    if text_content:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph
                        from reportlab.lib.styles import getSampleStyleSheet
                        
                        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
                        styles = getSampleStyleSheet()
                        elements = []
                        for line in text_content.split('\n'):
                            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            elements.append(Paragraph(line or '&nbsp;', styles['Normal']))
                        doc.build(elements)
                        result = str(output_path)
            
            # Boshqa formatlarga konvertatsiya
            elif output_format == 'txt':
                text_content = extract_text_content(input_path)
                if text_content:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                    result = str(output_path)
            
            elif output_format == 'docx':
                text_content = extract_text_content(input_path)
                if text_content:
                    new_doc = docx.Document()
                    for para in text_content.split('\n'):
                        new_doc.add_paragraph(para)
                    new_doc.save(str(output_path))
                    result = str(output_path)
            
            elif output_format == 'xlsx':
                text_content = extract_text_content(input_path)
                if text_content:
                    import pandas as pd
                    lines = text_content.split('\n')
                    df = pd.DataFrame({'Matn': lines})
                    df.to_excel(str(output_path), index=False)
                    result = str(output_path)
            
            if result and os.path.exists(result):
                return JsonResponse({
                    'success': True,
                    'filename': os.path.basename(result),
                    'download_url': f'/download/{os.path.basename(result)}/'
                })
            return JsonResponse({'success': False, 'error': 'Konvertatsiya muvaffaqiyatsiz. LibreOffice o\'rnatilganligini tekshiring.'})
        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'POST kerak'})


def extract_text_content(input_path):
    """Fayldan matn chiqarib olish"""
    input_path = Path(input_path)
    ext = input_path.suffix.lower()
    text_content = None
    
    try:
        import textract
        text_content = textract.process(str(input_path)).decode('utf-8')
    except:
        if ext in ['.txt', '.html', '.htm', '.css', '.js', '.py', '.json', '.xml', '.md']:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
        elif ext in ['.docx']:
            try:
                doc = docx.Document(str(input_path))
                text_content = '\n'.join([para.text for para in doc.paragraphs])
            except:
                pass
        elif ext in ['.xlsx', '.xls']:
            try:
                import pandas as pd
                df = pd.read_excel(str(input_path))
                text_content = df.to_string()
            except:
                pass
    
    return text_content


def convert_to_html_for_view(input_path):
    """LibreOffice orqali faylni HTML ga aylantirish (ko'rish uchun)"""
    input_path = Path(input_path)
    ext = input_path.suffix.lower()
    
    # Agar oddiy matnli fayl bo'lsa
    if ext in ['.txt', '.css', '.js', '.py', '.json', '.xml', '.md']:
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            # HTML ga o'tkazish (kodlarni ko'rsatish uchun)
            import html as html_module
            escaped = html_module.escape(content)
            return f'<pre style="white-space: pre-wrap; word-wrap: break-word; font-family: Consolas, Monaco, monospace;">{escaped}</pre>', 'html'
        except:
            return None, None
    
    # Agar HTML fayl bo'lsa
    if ext in ['.html', '.htm']:
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(), 'html'
        except:
            return None, None
    
    # LibreOffice orqali HTML ga aylantirish
    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    
    if not os.path.exists(libreoffice_path):
        return None, None
    
    if ext not in ['.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp', '.rtf']:
        return None, None
    
    try:
        import tempfile
        import subprocess
        import shutil
        
        # Vaqtinchalik papka yaratish
        temp_dir = tempfile.mkdtemp()
        
        # Faylni vaqtinchalik papkaga nusxalash
        temp_input = os.path.join(temp_dir, input_path.name)
        shutil.copy2(str(input_path), temp_input)
        
        # LibreOffice orqali HTML ga aylantirish
        cmd = [
            libreoffice_path,
            '--headless',
            '--convert-to', 'html',
            '--outdir', temp_dir,
            temp_input
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        # HTML faylni topish
        html_filename = input_path.stem + '.html'
        html_path = os.path.join(temp_dir, html_filename)
        
        if os.path.exists(html_path):
            with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            # Vaqtinchalik papkani o'chirish
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return html_content, 'html'
        else:
            shutil.rmtree(temp_dir, ignore_errors=True)
            return None, None
    except Exception as e:
        print(f"LibreOffice HTML konvertatsiya xatosi: {e}")
        return None, None


def edit_file_page(request):
    """Fayl tahrirlash sahifasi"""
    filepath = request.GET.get('filepath', '')
    filename = request.GET.get('filename', '')
    return render(request, 'blog/edit.html', {'filepath': filepath, 'filename': filename})


@csrf_exempt
def get_file_stats(request):
    """Fayl statistikasi va qidiruv"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            filepath = data.get('filepath')
            search_text = data.get('search', '')
            
            if not filepath or not os.path.exists(filepath):
                return JsonResponse({'success': False, 'error': 'Fayl topilmadi'})
            
            # Fayl matnini olish (statistika uchun)
            text_content = extract_text_content(Path(filepath))
            
            if not text_content:
                text_content = ''
            
            # LibreOffice orqali HTML ko'rinishini olish
            html_content, content_type = convert_to_html_for_view(Path(filepath))
            
            # Statistika hisoblash
            import re
            
            # So'zlar soni
            words = re.findall(r'\b\w+\b', text_content)
            word_count = len(words)
            
            # Gaplar soni (. ! ? bilan tugaydigan)
            sentences = re.split(r'[.!?]+', text_content)
            sentence_count = len([s for s in sentences if s.strip()])
            
            # Paragraflar soni
            paragraphs = [p for p in text_content.split('\n\n') if p.strip()]
            paragraph_count = len(paragraphs)
            
            # Belgilar soni
            char_count = len(text_content)
            char_count_no_spaces = len(text_content.replace(' ', '').replace('\n', ''))
            
            # Qidiruv
            search_results = []
            if search_text:
                pattern = re.compile(re.escape(search_text), re.IGNORECASE)
                matches = list(pattern.finditer(text_content))
                
                for match in matches:
                    start = max(0, match.start() - 50)
                    end = min(len(text_content), match.end() + 50)
                    context = text_content[start:end]
                    
                    # Topilgan so'zni ajratib ko'rsatish
                    highlighted = context.replace(match.group(), f'<mark>{match.group()}</mark>')
                    
                    search_results.append({
                        'position': match.start(),
                        'context': highlighted
                    })
            
            return JsonResponse({
                'success': True,
                'content': text_content,
                'html_content': html_content,
                'content_type': content_type,
                'stats': {
                    'word_count': word_count,
                    'sentence_count': sentence_count,
                    'paragraph_count': paragraph_count,
                    'char_count': char_count,
                    'char_count_no_spaces': char_count_no_spaces
                },
                'search_results': search_results,
                'search_count': len(search_results)
            })
        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'POST kerak'})


from django.http import FileResponse
def download_file(request, filename):
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        file_path = TEMP_DIR / filename
    if file_path.exists():
        response = FileResponse(open(file_path, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    return HttpResponse("Fayl topilmadi", status=404)


@csrf_exempt
def convert_folder_to_pdf_zip(request):
    """Papka ichidagi barcha fayllarni PDF ga aylantirib, ZIP arxivga joylashtirish"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            files_data = data.get('files', [])
            folder_name = data.get('folder_name', 'papka')
            
            if not files_data:
                return JsonResponse({'success': False, 'error': 'Fayllar topilmadi'})
            
            # PDF fayllarni saqlash uchun vaqtinchalik papka
            import tempfile
            import shutil
            temp_pdf_dir = Path(tempfile.mkdtemp())
            
            converted_pdfs = []
            
            for file_info in files_data:
                filepath = file_info.get('path')
                if not filepath or not os.path.exists(filepath):
                    continue
                
                input_path = Path(filepath)
                ext = input_path.suffix.lower()
                pdf_output = temp_pdf_dir / f"{input_path.stem}.pdf"
                
                # LibreOffice orqali konvertatsiya
                libre_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    r"C:\Program Files\LibreOffice\program\soffice.com",
                    "/usr/bin/libreoffice",
                    "/usr/bin/soffice"
                ]
                
                libre_path = None
                for path in libre_paths:
                    if os.path.exists(path):
                        libre_path = path
                        break
                
                if libre_path and ext in ['.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp', '.rtf', '.txt', '.html', '.htm']:
                    try:
                        import subprocess
                        result_process = subprocess.run([
                            libre_path,
                            '--headless',
                            '--convert-to', 'pdf',
                            '--outdir', str(temp_pdf_dir),
                            str(input_path)
                        ], capture_output=True, timeout=60)
                        
                        if pdf_output.exists():
                            converted_pdfs.append(pdf_output)
                    except Exception as e:
                        print(f"LibreOffice xatolik: {e}")
                
                elif ext == '.pdf':
                    # Allaqachon PDF - nusxa olish
                    shutil.copy(str(input_path), str(pdf_output))
                    converted_pdfs.append(pdf_output)
                
                else:
                    # Boshqa usullar bilan konvertatsiya
                    try:
                        result = convert_any_to_pdf(str(input_path), str(pdf_output))
                        if result and os.path.exists(result):
                            converted_pdfs.append(Path(result))
                    except:
                        pass
            
            if not converted_pdfs:
                shutil.rmtree(temp_pdf_dir, ignore_errors=True)
                return JsonResponse({'success': False, 'error': 'Hech bir fayl konvertatsiya qilinmadi'})
            
            # ZIP arxiv yaratish
            zip_filename = f"{folder_name}_pdf.zip"
            zip_path = OUTPUT_DIR / zip_filename
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for pdf_file in converted_pdfs:
                    zipf.write(pdf_file, pdf_file.name)
            
            # Vaqtinchalik papkani o'chirish
            shutil.rmtree(temp_pdf_dir, ignore_errors=True)
            
            return JsonResponse({
                'success': True,
                'filename': zip_filename,
                'download_url': f'/download/{zip_filename}/',
                'converted_count': len(converted_pdfs)
            })
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'POST kerak'})


@csrf_exempt
def delete_file(request):
    """Faylni serverdan va bazadan o'chirish"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
            filepath = data.get('filepath')
            file_id = data.get('file_id')
            
            if not filepath:
                return JsonResponse({'success': False, 'error': 'Fayl yo\'li ko\'rsatilmagan'})
            
            # Bazadan o'chirish
            if file_id:
                try:
                    file_obj = File.objects.get(id=file_id)
                    file_obj.delete()
                except File.DoesNotExist:
                    pass
            else:
                # Filepath bo'yicha bazadan topib o'chirish
                for f in File.objects.all():
                    if f.file and f.file.path == filepath:
                        f.delete()
                        break
            
            # Faylni diskdan o'chirish
            if os.path.exists(filepath):
                os.remove(filepath)
                return JsonResponse({'success': True, 'message': 'Fayl o\'chirildi'})
            else:
                return JsonResponse({'success': True, 'message': 'Fayl bazadan o\'chirildi'})
        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'POST kerak'})


@login_required
def yuklangan_fayllar(request):
    fs = FileSystemStorage()
    if request.method == 'POST':
        if 'delete' in request.POST:
            selected_files = request.POST.getlist('selected_files')
            File.objects.filter(user=request.user, file__in=selected_files).delete()
            return redirect('yuklangan_fayllar')
        elif 'download' in request.POST:
            selected_files = request.POST.getlist('selected_files')
            user_files = File.objects.filter(user=request.user, file__in=selected_files)
            if user_files:
                buffer = BytesIO()
                with zipfile.ZipFile(buffer, 'w') as zip_file:
                    for file_obj in user_files:
                        file_path = file_obj.file.path
                        if os.path.exists(file_path):
                            zip_file.write(file_path, file_obj.file.name)
                buffer.seek(0)
                response = HttpResponse(buffer.getvalue(), content_type='application/zip')
                response['Content-Disposition'] = 'attachment; filename="selected_files.zip"'
                return response
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.heic', '.svg', '.raw']
    files = [f.file.name for f in File.objects.filter(user=request.user).exclude(file__endswith=tuple(image_extensions))]
    return render(request, "blog/malumot.html", {"files": files, "filename": "Yuklangan fayllar"})


@login_required
def image_list(request):
    if request.method == 'POST':
        # Rasm yuklash
        if 'upload_image' in request.POST:
            if request.FILES.get('image'):
                title = request.POST.get('title', '').strip()
                image_file = request.FILES['image']
                new_image = Image(
                    user=request.user,
                    title=title if title else image_file.name,
                    image=image_file
                )
                new_image.save()
                messages.success(request, 'Rasm muvaffaqiyatli yuklandi!')
            return redirect('image_list')
        
        # Rasmlarni o'chirish
        elif 'delete' in request.POST:
            selected_images = request.POST.getlist('selected_images')
            for image_id in selected_images:
                try:
                    image = Image.objects.get(id=image_id)
                    # Delete the file
                    if image.image:
                        image.image.delete(save=False)
                    # Delete the model instance
                    image.delete()
                except Image.DoesNotExist:
                    pass
            return redirect('image_list')
        elif 'convert' in request.POST:
            to_format = request.POST['convert']
            selected_ids = request.POST.getlist('selected_images')
            if not selected_ids:
                messages.error(request, 'Iltimos, konvertatsiya qilish uchun rasm tanlang')
                return redirect('image_list')
            converted_files = []
            for id in selected_ids:
                try:
                    image = Image.objects.get(id=id, user=request.user)
                    filename = image.image.name
                    ext = filename.split('.')[-1].lower()
                    
                    # Fayl yo'lini to'g'ri olish
                    file_path = image.image.path
                    
                    if to_format == 'pdf':
                        try:
                            img = PILImage.open(file_path)
                            # RGBA rasmlarni RGB ga o'tkazish (PDF uchun shart)
                            if img.mode in ('RGBA', 'P', 'LA'):
                                img = img.convert('RGB')
                            
                            # Pillow bilan to'g'ridan-to'g'ri PDF ga saqlash
                            buffer = BytesIO()
                            img.save(buffer, format='PDF', resolution=100.0)
                            buffer.seek(0)
                            converted_files.append((f"{os.path.splitext(os.path.basename(filename))[0]}.pdf", buffer.getvalue()))
                        except Exception as e:
                            messages.error(request, f'PDF yaratishda xatolik: {str(e)}')
                            continue
                    elif to_format in ['png', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
                        try:
                            img = PILImage.open(file_path)
                            buffer = BytesIO()
                            if to_format == 'jpeg':
                                if img.mode in ("RGBA", "P"):
                                    img = img.convert("RGB")
                                img.save(buffer, format='JPEG')
                                new_ext = 'jpg'
                            elif to_format == 'png':
                                img.save(buffer, format='PNG')
                                new_ext = 'png'
                            elif to_format == 'gif':
                                img.save(buffer, format='GIF')
                                new_ext = 'gif'
                            elif to_format == 'bmp':
                                img.save(buffer, format='BMP')
                                new_ext = 'bmp'
                            elif to_format == 'tiff':
                                img.save(buffer, format='TIFF')
                                new_ext = 'tiff'
                            elif to_format == 'webp':
                                img.save(buffer, format='WEBP')
                                new_ext = 'webp'
                            buffer.seek(0)
                            converted_files.append((f"{os.path.splitext(os.path.basename(filename))[0]}.{new_ext}", buffer.getvalue()))
                        except Exception as e:
                            messages.error(request, f'Konvertatsiya xatoligi: {str(e)}')
                            continue
                except Image.DoesNotExist:
                    messages.error(request, f'Rasm topilmadi: ID {id}')
                    continue
                except Exception as e:
                    messages.error(request, f'Xatolik: {str(e)}')
                    continue
            
            if len(converted_files) == 1:
                name, data = converted_files[0]
                # Fayl nomini ASCII ga o'zgartirish
                import re
                import time
                # Faqat ASCII belgilarni qoldirish
                ascii_name = re.sub(r'[^\x00-\x7F]+', '', name)
                # Agar nom bo'sh qolsa, yangi nom berish
                if not ascii_name or ascii_name == '.pdf':
                    ascii_name = f"converted_{int(time.time())}.pdf"
                # Bo'sh joylarni olib tashlash
                ascii_name = ascii_name.replace(' ', '_').strip('_-')
                
                response = HttpResponse(data, content_type='application/octet-stream')
                response['Content-Disposition'] = f'attachment; filename="{ascii_name}"'
                response['Content-Length'] = len(data)
                return response
            elif len(converted_files) > 1:
                buffer = BytesIO()
                with zipfile.ZipFile(buffer, 'w') as zf:
                    for name, data in converted_files:
                        zf.writestr(name, data)
                buffer.seek(0)
                response = HttpResponse(buffer.getvalue(), content_type='application/zip')
                response['Content-Disposition'] = 'attachment; filename="converted_images.zip"'
                return response
            else:
                messages.error(request, 'Konvertatsiya qilish uchun rasm tanlanmadi yoki xatolik yuz berdi')
                return redirect('image_list')
    images = Image.objects.filter(user=request.user)
    for image in images:
        image.extension = image.image.name.split('.')[-1].upper() if '.' in image.image.name else 'NOMA\'LUM'
    return render(request, 'blog/image_list.html', {'images': images})


def convert_file(request, filename, to_format):
    fs = FileSystemStorage()
    file_path = fs.path(filename)
    if not os.path.exists(file_path):
        raise Http404("Fayl topilmadi")

    ext = filename.split('.')[-1].lower()

    if to_format == 'txt':
        # Extract text and return as TXT
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"

        response = HttpResponse(text, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
        return response

    elif to_format == 'pdf':
        # Word, Excel, PowerPoint fayllar uchun LibreOffice ishlatish
        if ext.lower() in ['doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx', 'odt', 'ods', 'odp']:
            try:
                # LibreOffice bilan konvertatsiya
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    temp_pdf_path = temp_file.name
                
                result = convert_to_pdf_with_libreoffice(file_path, temp_pdf_path)
                
                if result and os.path.exists(temp_pdf_path):
                    with open(temp_pdf_path, 'rb') as f:
                        pdf_content = f.read()
                    
                    # Vaqtinchalik faylni o'chirish
                    try:
                        os.unlink(temp_pdf_path)
                    except:
                        pass
                    
                    response = HttpResponse(pdf_content, content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{os.path.splitext(filename)[0]}.pdf"'
                    return response
                else:
                    raise Exception("LibreOffice konvertatsiya muvaffaqiyatsiz")
                
            except Exception as e:
                return HttpResponse(f"PDF yaratishda xatolik: {str(e)}", status=500)
        
        # Boshqa formatlar uchun (TXT, rasm va h.k.) - reportlab bilan
        try:
            # Agar fayl allaqachon PDF bo'lsa, to'g'ridan-to'g'ri qaytarish
            if ext.lower() == 'pdf':
                with open(file_path, 'rb') as f:
                    response = HttpResponse(f.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
            
            # TXT va boshqa matnli fayllar uchun reportlab bilan PDF yaratish
            text = ""
            try:
                if ext == "txt":
                    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                        text = f.read()
                else:
                    text = "Bu format qo'llab-quvvatlanmaydi."
            except Exception as e:
                text = f"Xatolik: {str(e)}"

            # Create PDF
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)
            p.setFont("Helvetica", 12)
            lines = text.split('\n')
            y = 750
            for line in lines:
                if y < 50:
                    p.showPage()
                    p.setFont("Helvetica", 12)
                    y = 750
                # Limit line length to fit page
                while line:
                    part = line[:80]
                    line = line[80:]
                    p.drawString(50, y, part)
                    y -= 15
                    if y < 50:
                        p.showPage()
                        p.setFont("Helvetica", 12)
                        y = 750
                p.save()
                buffer.seek(0)

                response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
                return response
                
        except Exception as e:
            # Xato bo'lsa, eski usulni ishlatish
            return HttpResponse(f"PDF yaratishda xatolik: {str(e)}", status=500)

    elif to_format == 'doc':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        doc = docx.Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/msword')
        response['Content-Disposition'] = f'attachment; filename="{filename}.doc"'
        return response

    elif to_format == 'docx':
        # PDF fayllar uchun CloudConvert ishlatish
        if ext.lower() == 'pdf':
            try:
                import requests
                
                # CloudConvert API bilan konvertatsiya
                job = cloudconvert.Job.create(payload={
                    "tasks": {
                        "import-my-file": {
                            "operation": "import/upload"
                        },
                        "convert-my-file": {
                            "operation": "convert",
                            "input": "import-my-file",
                            "output_format": "docx"
                        },
                        "export-my-file": {
                            "operation": "export/url",
                            "input": "convert-my-file"
                        }
                    }
                })
                
                # Faylni yuklash
                upload_task = None
                for task in job["tasks"]:
                    if task["name"] == "import-my-file":
                        upload_task = task
                        break
                
                if upload_task:
                    upload_task = cloudconvert.Task.find(id=upload_task["id"])
                    with open(file_path, 'rb') as f:
                        cloudconvert.Task.upload(file_name=filename, task=upload_task, file=f)
                    
                    # Job tugashini kutish
                    job = cloudconvert.Job.wait(id=job["id"])
                    
                    # Natijani yuklab olish
                    export_task = None
                    for task in job["tasks"]:
                        if task["name"] == "export-my-file" and task["status"] == "finished":
                            export_task = task
                            break
                    
                    if export_task and export_task.get("result") and export_task["result"].get("files"):
                        file_url = export_task["result"]["files"][0]["url"]
                        docx_response = requests.get(file_url)
                        
                        response = HttpResponse(docx_response.content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename="{os.path.splitext(filename)[0]}.docx"'
                        return response
                
                raise Exception("CloudConvert konvertatsiya muvaffaqiyatsiz")
                
            except Exception as e:
                # Fallback - oddiy usul bilan
                pass
        
        # Boshqa formatlar yoki fallback uchun eski usul
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        doc = docx.Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    elif to_format == 'rtf':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        rtf_content = "{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0 Times New Roman;}}\\f0\\fs24\n"
        for line in text.split('\n'):
            rtf_content += line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + "\\par\n"
        rtf_content += "}"
        response = HttpResponse(rtf_content, content_type='application/rtf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.rtf"'
        return response

    elif to_format == 'odt':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        doc = OpenDocumentText()
        p = odf_text.P()
        teletype.addTextToElement(p, text)
        doc.text.addElement(p)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.oasis.opendocument.text')
        response['Content-Disposition'] = f'attachment; filename="{filename}.odt"'
        return response

    elif to_format == 'html':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        html_content = f"<html><head><title>{filename}</title></head><body><pre>{text}</pre></body></html>"
        response = HttpResponse(html_content, content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="{filename}.html"'
        return response

    elif to_format == 'xml':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        xml_content = f"<root><text><![CDATA[{text}]]></text></root>"
        response = HttpResponse(xml_content, content_type='application/xml')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xml"'
        return response

    elif to_format == 'epub':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        book = epub.EpubBook()
        book.set_identifier(filename)
        book.set_title(filename)
        book.set_language('uz')
        c1 = epub.EpubHtml(title='Text', file_name='chap1.xhtml', lang='uz')
        c1.content = f"<h1>{filename}</h1><pre>{text}</pre>"
        book.add_item(c1)
        book.toc = (c1,)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav', c1]
        buffer = BytesIO()
        epub.write_epub(buffer, book)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/epub+zip')
        response['Content-Disposition'] = f'attachment; filename="{filename}.epub"'
        return response

    elif to_format == 'mobi':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        book = epub.EpubBook()
        book.set_identifier(filename)
        book.set_title(filename)
        book.set_language('uz')
        c1 = epub.EpubHtml(title='Text', file_name='chap1.xhtml', lang='uz')
        c1.content = f"<h1>{filename}</h1><pre>{text}</pre>"
        book.add_item(c1)
        book.toc = (c1,)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav', c1]
        buffer = BytesIO()
        epub.write_epub(buffer, book)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/x-mobipocket-ebook')
        response['Content-Disposition'] = f'attachment; filename="{filename}.mobi"'
        return response

    elif to_format in ['xlsx', 'xlsm', 'xltx', 'xltm']:
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        wb = openpyxl.Workbook()
        ws = wb.active
        for i, line in enumerate(text.split('\n'), 1):
            ws.cell(row=i, column=1).value = line
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.{to_format}"'
        return response

    elif to_format == 'xls':
        text = ""
        try:
            if ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            elif ext in ["doc", "docx"]:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
            else:
                text = "Bu format qo'llab-quvvatlanmaydi."
        except Exception as e:
            text = f"Xatolik: {str(e)}"
        wb = xlwt.Workbook()
        ws = wb.add_sheet('Sheet1')
        for i, line in enumerate(text.split('\n')):
            ws.write(i, 0, line)
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.ms-excel')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xls"'
        return response

    elif to_format in ['xlsb', 'xlt']:
        raise Http404("Bu Excel format hozircha qo'llab-quvvatlanmaydi")

    elif to_format in ['jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp', 'heic', 'svg', 'raw']:
        try:
            img = PILImage.open(file_path)
            buffer = BytesIO()
            if to_format == 'jpeg':
                # Convert to RGB if necessary
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(buffer, format='JPEG')
                content_type = 'image/jpeg'
                ext = 'jpg'
            elif to_format == 'png':
                img.save(buffer, format='PNG')
                content_type = 'image/png'
                ext = 'png'
            elif to_format == 'gif':
                img.save(buffer, format='GIF')
                content_type = 'image/gif'
                ext = 'gif'
            elif to_format == 'bmp':
                img.save(buffer, format='BMP')
                content_type = 'image/bmp'
                ext = 'bmp'
            elif to_format == 'tiff':
                img.save(buffer, format='TIFF')
                content_type = 'image/tiff'
                ext = 'tiff'
            elif to_format == 'webp':
                img.save(buffer, format='WEBP')
                content_type = 'image/webp'
                ext = 'webp'
            elif to_format == 'heic':
                img.save(buffer, format='HEIC')
                content_type = 'image/heic'
                ext = 'heic'
            elif to_format == 'svg':
                # SVG is vector, but if Pillow can save as SVG, or convert
                # Pillow doesn't save SVG, so perhaps convert to PNG or something
                # For now, save as PNG since SVG output not supported
                img.save(buffer, format='PNG')
                content_type = 'image/png'
                ext = 'png'  # Since we save as PNG
            elif to_format == 'raw':
                # RAW is not a standard format, perhaps save as TIFF or something
                img.save(buffer, format='TIFF')
                content_type = 'image/tiff'
                ext = 'tiff'
            buffer.seek(0)
            base_name = os.path.splitext(filename)[0]
            response = HttpResponse(buffer.getvalue(), content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{base_name}.{ext}"'
            return response
        except Exception as e:
            return HttpResponse(f"Xato: {e}", status=400)

    else:
        raise Http404("Format qo'llab-quvvatlanmaydi")


def feedback_list(request):
    password = request.GET.get('password')
    if password != 'Akobir05':
        return HttpResponse("Parol noto'g'ri", status=403)
    feedbacks = Feedback.objects.all().order_by('-created_at')
    return render(request, 'blog/feedback_list.html', {'feedbacks': feedbacks})


def submit_feedback(request):
    if request.method == 'POST':
        name = request.POST.get('name', 'Anonim' if not request.user.is_authenticated else request.user.username)
        message = request.POST.get('message', '')
        if message:
            Feedback.objects.create(user=request.user if request.user.is_authenticated else None, name=name, message=message)
        return redirect('boshlash')
    return redirect('boshlash')


def register(request):
    from django.core.mail import send_mail
    from django.contrib.auth.hashers import make_password
    from .models import EmailVerification
    import re
    
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip().lower()
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        password1 = request.POST.get('password1', '')
        password2 = request.POST.get('password2', '')
        
        errors = []
        
        # Ism va familiya tekshirish
        if not first_name:
            errors.append('Ismingizni kiriting')
        elif len(first_name) < 2:
            errors.append('Ism kamida 2 ta harfdan iborat bo\'lishi kerak')
            
        if not last_name:
            errors.append('Familiyangizni kiriting')
        elif len(last_name) < 2:
            errors.append('Familiya kamida 2 ta harfdan iborat bo\'lishi kerak')
        
        # Email tekshirish
        if not email:
            errors.append('Email manzilingizni kiriting')
        elif not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
            errors.append('Noto\'g\'ri email formati')
        elif User.objects.filter(email__iexact=email).exists():
            errors.append('Bu email allaqachon ro\'yxatdan o\'tgan')
        
        if not username:
            errors.append('Foydalanuvchi nomini kiriting')
        elif len(username) < 3:
            errors.append('Foydalanuvchi nomi kamida 3 ta belgidan iborat bo\'lishi kerak')
        elif User.objects.filter(username__iexact=username).exists():
            errors.append('Bu foydalanuvchi nomi band. Boshqa nom tanlang')
            
        if not password1:
            errors.append('Parolni kiriting')
        if not password2:
            errors.append('Parolni tasdiqlang')
        
        if password1 and password2 and password1 != password2:
            errors.append('Parollar bir xil emas')
        
        if password1 and len(password1) < 4:
            errors.append('Parol kamida 4 ta belgidan iborat bo\'lishi kerak')
        
        if errors:
            return render(request, 'registration/register.html', {
                'errors': errors,
                'username': username,
                'email': email,
                'first_name': first_name,
                'last_name': last_name
            })
        
        try:
            # Parolni hash qilish
            hashed_password = make_password(password1)
            
            # Email tasdiqlash yaratish
            verification = EmailVerification.create_verification(
                email=email,
                username=username,
                password=hashed_password,
                first_name=first_name,
                last_name=last_name
            )
            
            # Email yuborish
            try:
                send_mail(
                    subject='Cloudstore - Email tasdiqlash kodi',
                    message=f'''Assalomu alaykum, {first_name}!

Cloudstore'da ro'yxatdan o'tish uchun tasdiqlash kodingiz:

🔐 {verification.code}

Bu kod 15 daqiqa ichida amal qiladi.

Agar siz ro'yxatdan o'tmagan bo'lsangiz, bu xabarni e'tiborsiz qoldiring.

Hurmat bilan,
Cloudstore jamoasi
''',
                    from_email=None,  # DEFAULT_FROM_EMAIL ishlatiladi
                    recipient_list=[email],
                    fail_silently=False,
                )
            except Exception as e:
                # Email yuborishda xatolik bo'lsa
                verification.delete()
                errors.append(f'Email yuborishda xatolik: {str(e)}')
                return render(request, 'registration/register.html', {
                    'errors': errors,
                    'username': username,
                    'email': email,
                    'first_name': first_name,
                    'last_name': last_name
                })
            
            # Tasdiqlash sahifasiga yo'naltirish
            return redirect('verify_email', verification_id=verification.id)
            
        except Exception as e:
            errors.append(f'Xatolik: {str(e)}')
            return render(request, 'registration/register.html', {
                'errors': errors,
                'username': username,
                'email': email,
                'first_name': first_name,
                'last_name': last_name
            })
    
    return render(request, 'registration/register.html')


def verify_email(request, verification_id):
    """Email tasdiqlash sahifasi"""
    from .models import EmailVerification
    from django.contrib.auth.hashers import check_password
    
    try:
        verification = EmailVerification.objects.get(id=verification_id, is_verified=False)
    except EmailVerification.DoesNotExist:
        messages.error(request, 'Tasdiqlash so\'rovi topilmadi yoki allaqachon tasdiqlangan.')
        return redirect('register')
    
    # Muddati tugagan tekshirish
    if verification.is_expired():
        verification.delete()
        messages.error(request, 'Tasdiqlash kodi muddati tugagan. Qaytadan ro\'yxatdan o\'ting.')
        return redirect('register')
    
    if request.method == 'POST':
        code = request.POST.get('code', '').strip()
        
        if not code:
            return render(request, 'registration/verify_email.html', {
                'verification': verification,
                'error': 'Tasdiqlash kodini kiriting',
                'email': verification.email
            })
        
        success, message = verification.verify(code)
        
        if success:
            # Foydalanuvchi yaratish
            try:
                user = User.objects.create(
                    username=verification.username,
                    email=verification.email,
                    first_name=verification.first_name,
                    last_name=verification.last_name,
                    password=verification.password  # Allaqachon hash qilingan
                )
                
                # Avtomatik kirish
                login(request, user)
                messages.success(request, f"Xush kelibsiz, {verification.first_name}! Email muvaffaqiyatli tasdiqlandi. 🎉")
                
                # Tasdiqlash yozuvini o'chirish
                verification.delete()
                
                return redirect('boshlash')
                
            except Exception as e:
                return render(request, 'registration/verify_email.html', {
                    'verification': verification,
                    'error': f'Foydalanuvchi yaratishda xatolik: {str(e)}',
                    'email': verification.email
                })
        else:
            return render(request, 'registration/verify_email.html', {
                'verification': verification,
                'error': message,
                'email': verification.email
            })
    
    return render(request, 'registration/verify_email.html', {
        'verification': verification,
        'email': verification.email
    })


def resend_verification_code(request, verification_id):
    """Tasdiqlash kodini qayta yuborish"""
    from django.core.mail import send_mail
    from .models import EmailVerification
    from django.utils import timezone
    from datetime import timedelta
    
    try:
        verification = EmailVerification.objects.get(id=verification_id, is_verified=False)
    except EmailVerification.DoesNotExist:
        messages.error(request, 'Tasdiqlash so\'rovi topilmadi.')
        return redirect('register')
    
    # Yangi kod yaratish
    verification.code = EmailVerification.generate_code()
    verification.expires_at = timezone.now() + timedelta(minutes=15)
    verification.attempts = 0
    verification.save()
    
    # Email yuborish
    try:
        send_mail(
            subject='Cloudstore - Yangi tasdiqlash kodi',
            message=f'''Assalomu alaykum, {verification.first_name}!

Yangi tasdiqlash kodingiz:

🔐 {verification.code}

Bu kod 15 daqiqa ichida amal qiladi.

Hurmat bilan,
Cloudstore jamoasi
''',
            from_email=None,
            recipient_list=[verification.email],
            fail_silently=False,
        )
        messages.success(request, 'Yangi tasdiqlash kodi yuborildi!')
    except Exception as e:
        messages.error(request, f'Email yuborishda xatolik: {str(e)}')
    
    return redirect('verify_email', verification_id=verification.id)


def user_login(request):
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '')
        
        # Bo'sh maydonlarni tekshirish
        if not username or not password:
            error = 'Foydalanuvchi nomi va parolni kiriting'
            return render(request, 'registration/login.html', {'error': error})
        
        # Foydalanuvchini bazadan qidirish
        try:
            user_exists = User.objects.filter(username=username).exists()
            if not user_exists:
                # Username bo'yicha topilmasa, case-insensitive qidirish
                user_exists = User.objects.filter(username__iexact=username).exists()
                if user_exists:
                    # To'g'ri username ni olish
                    correct_user = User.objects.get(username__iexact=username)
                    username = correct_user.username
        except Exception as e:
            error = f'Xatolik yuz berdi: {str(e)}'
            return render(request, 'registration/login.html', {'error': error})
        
        # Autentifikatsiya
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            if user.is_active:
                login(request, user)
                messages.success(request, f'Xush kelibsiz, {user.username}!')
                return redirect('boshlash')
            else:
                error = 'Foydalanuvchi hisobi faol emas'
                return render(request, 'registration/login.html', {'error': error})
        else:
            # Foydalanuvchi topilmadi yoki parol noto'g'ri
            if User.objects.filter(username=username).exists():
                error = 'Parol noto\'g\'ri. Qayta urinib ko\'ring'
            else:
                error = 'Bunday foydalanuvchi yo\'q. Ro\'yxatdan o\'ting'
            return render(request, 'registration/login.html', {'error': error, 'username': username})
    
    return render(request, 'registration/login.html')


# ============= AI CHAT FUNKSIYALARI =============

@login_required
def ai_chat(request):
    """AI chat sahifasi"""
    # Foydalanuvchining fayllarini olish
    files = File.objects.filter(user=request.user).order_by('-uploaded_at')
    file_list = []
    for f in files:
        if f.file and os.path.exists(f.file.path):
            file_list.append({
                'id': f.id,
                'name': os.path.basename(f.file.name),
                'path': f.file.path
            })
    return render(request, 'blog/ai_chat.html', {'files': file_list})


@csrf_exempt
def get_file_content(request):
    """Fayl kontentini olish"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST kerak'})
    
    try:
        data = json.loads(request.body)
        file_id = data.get('file_id')
        
        file_obj = File.objects.get(id=file_id)
        file_path = file_obj.file.path
        filename = os.path.basename(file_obj.file.name)
        ext = filename.split('.')[-1].lower()
        
        content = ""
        
        # Fayl turini aniqlash va kontentni o'qish
        if ext == 'txt':
            encodings = ['utf-8', 'cp1251', 'latin-1', 'utf-16']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except:
                    continue
        
        elif ext in ['doc', 'docx']:
            try:
                doc = docx.Document(file_path)
                content = '\n'.join([p.text for p in doc.paragraphs])
            except Exception as e:
                content = f"DOCX o'qishda xatolik: {str(e)}"
        
        elif ext == 'pdf':
            try:
                reader = PdfReader(file_path)
                content = '\n'.join([page.extract_text() or '' for page in reader.pages])
            except Exception as e:
                content = f"PDF o'qishda xatolik: {str(e)}"
        
        elif ext in ['csv']:
            import csv
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    rows = list(reader)
                    content = '\n'.join([', '.join(row) for row in rows[:100]])  # Birinchi 100 qator
            except:
                content = "CSV o'qishda xatolik"
        
        elif ext in ['json']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except:
                content = "JSON o'qishda xatolik"
        
        elif ext in ['html', 'htm', 'xml']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except:
                content = "HTML/XML o'qishda xatolik"
        
        elif ext in ['py', 'js', 'css', 'java', 'cpp', 'c', 'h']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except:
                content = "Kod faylini o'qishda xatolik"
        
        else:
            content = "Bu fayl turi qo'llab-quvvatlanmaydi"
        
        # Kontentni 10000 belgiga cheklash
        if len(content) > 10000:
            content = content[:10000] + "\n\n... (qolgan qism qisqartirildi)"
        
        return JsonResponse({
            'success': True,
            'filename': filename,
            'content': content
        })
        
    except File.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Fayl topilmadi'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@csrf_exempt
def ai_ask(request):
    """AI ga savol berish"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST kerak'})
    
    try:
        data = json.loads(request.body)
        question = data.get('question', '')
        content = data.get('content', '')
        
        if not question:
            return JsonResponse({'success': False, 'error': 'Savol kiritilmadi'})
        
        if not content:
            return JsonResponse({'success': False, 'error': 'Fayl kontenti yo\'q'})
        
        # Oddiy AI - kalit so'zlar asosida javob
        answer = simple_ai_answer(question, content)
        
        return JsonResponse({
            'success': True,
            'answer': answer
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


def simple_ai_answer(question, content):
    """Groq AI - fayl kontenti asosida javob berish"""
    # Agar GROQ API kaliti sozlanmagan bo'lsa, foydalanuvchiga aniq xabar qaytaramiz
    if not groq_client:
        return "AI xizmati sozlanmagan: iltimos, GROQ_API_KEY muhit o'zgaruvchisini o'rnating."

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "Sen yordamchi AI assistantisan. Foydalanuvchi senga fayl kontentini beradi va u haqida savol so'raydi. Javobni o'zbek tilida, qisqa va aniq ber."
                },
                {
                    "role": "user",
                    "content": f"""Quyidagi fayl kontentiga asoslanib savolga javob ber:

=== FAYL KONTENTI ===
{content[:12000]}
=== FAYL KONTENTI TUGADI ===

Savol: {question}"""
                }
            ],
            temperature=0.7,
            max_tokens=1024
        )

        return response.choices[0].message.content
        
    except Exception as e:
        # Xatolik bo'lsa
        return f"AI xatolik: {str(e)}. Iltimos, keyinroq urinib ko'ring."


def kitob_yuklash(request):
    """Kitob yuklash sahifasi - foydalanuvchilar kitob qo'sha oladi"""
    from .models import Author, Book
    from django.contrib import messages
    
    if request.method == 'POST':
        # Formdan ma'lumotlarni olish
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        author_id = request.POST.get('author_id')
        new_author_name = request.POST.get('new_author_name', '').strip()
        published_year = request.POST.get('published_year', '').strip()
        book_file = request.FILES.get('book_file')
        
        # Validatsiya
        if not title:
            messages.error(request, "Kitob nomi kiritilishi shart!")
            return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
        
        if not book_file:
            messages.error(request, "Kitob fayli yuklanishi shart!")
            return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
        
        # Fayl formatini tekshirish - PDF qabul qilinmaydi!
        allowed_extensions = ['.docx', '.doc', '.txt']
        file_ext = os.path.splitext(book_file.name)[1].lower()
        
        if file_ext == '.pdf':
            messages.error(request, "❌ PDF format qabul qilinmaydi! Iltimos Word formatda (.docx, .doc) yuklang.")
            return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
        
        if file_ext not in allowed_extensions:
            messages.error(request, f"Faqat {', '.join(allowed_extensions)} formatdagi fayllar qabul qilinadi! PDF qabul qilinmaydi.")
            return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
        
        # Muallif tanlash yoki yangi yaratish
        author = None
        if new_author_name:
            # Yangi muallif yaratish
            author, created = Author.objects.get_or_create(name=new_author_name)
        elif author_id:
            try:
                author = Author.objects.get(id=author_id)
            except Author.DoesNotExist:
                messages.error(request, "Muallif topilmadi!")
                return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
        else:
            messages.error(request, "Muallif tanlang yoki yangi muallif nomini kiriting!")
            return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
        
        # Kitobni saqlash
        try:
            book = Book(
                title=title,
                description=description if description else '',
                author=author,
                file=book_file
            )
            
            if published_year:
                try:
                    book.year_written = int(published_year)
                except ValueError:
                    pass
            
            book.save()
            
            # Sahifalarni ajratib saqlash
            book.save_pages_from_file()
            
            messages.success(request, f"'{title}' kitobi muvaffaqiyatli yuklandi!")
            return render(request, 'blog/kitob_yuklash.html', {
                'authors': Author.objects.all(),
                'success': True,
                'book': book
            })
            
        except Exception as e:
            messages.error(request, f"Xatolik yuz berdi: {str(e)}")
            return render(request, 'blog/kitob_yuklash.html', {'authors': Author.objects.all()})
    
    # GET so'rovi - forma ko'rsatish
    authors = Author.objects.all().order_by('name')
    return render(request, 'blog/kitob_yuklash.html', {'authors': authors})


def get_authors(request):
    """Mualliflar ro'yxatini JSON formatida qaytarish"""
    from .models import Author
    
    authors = Author.objects.all().order_by('name')
    authors_list = [{'id': a.id, 'name': a.name} for a in authors]
    return JsonResponse({'authors': authors_list})


def ai_extract_book_info(content, filename):
    """AI yordamida kitob ma'lumotlarini aniqlash"""
    if not groq_client:
        return None
    
    try:
        # Kontentni qisqartirish (token limitiga moslashtirish)
        short_content = content[:3000] if content else ""
        
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": """Kitob ma'lumotlarini JSON formatda ber:
{"author": "Muallif ismi", "title": "Kitob nomi", "year": 2020, "description": "Qisqa tavsif"}
Agar ma'lumot topilmasa null yoz."""
                },
                {
                    "role": "user",
                    "content": f"""Fayl: {filename}
Matn: {short_content}

JSON formatda javob ber."""
                }
            ],
            temperature=0.3,
            max_tokens=300
        )
        
        result = response.choices[0].message.content.strip()
        
        # JSON ni ajratib olish
        import json
        import re
        
        # JSON qismini topish
        json_match = re.search(r'\{[^{}]*\}', result)
        if json_match:
            return json.loads(json_match.group())
        
        return None
        
    except Exception as e:
        print(f"AI xatolik: {e}")
        return None


@csrf_exempt
def upload_zip_books(request):
    """ZIP papkadan kitoblarni avtomatik yuklash"""
    from .models import Author, Book
    import zipfile
    import tempfile
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST so\'rov qabul qilinadi'}, status=400)
    
    zip_file = request.FILES.get('zip_file')
    if not zip_file:
        return JsonResponse({'error': 'ZIP fayl yuklanmadi'}, status=400)
    
    # ZIP ekanligini tekshirish
    if not zip_file.name.lower().endswith('.zip'):
        return JsonResponse({'error': 'Faqat ZIP formatdagi fayllar qabul qilinadi'}, status=400)
    
    results = []
    allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
    
    try:
        # Vaqtinchalik papkaga saqlash
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_path = os.path.join(temp_dir, 'uploaded.zip')
            
            with open(zip_path, 'wb') as f:
                for chunk in zip_file.chunks():
                    f.write(chunk)
            
            # ZIP ni ochish
            with zipfile.ZipFile(zip_path, 'r') as zf:
                file_list = zf.namelist()
                
                for file_name in file_list:
                    # Papkalarni o'tkazib yuborish
                    if file_name.endswith('/'):
                        continue
                    
                    # Fayl kengaytmasini tekshirish
                    ext = os.path.splitext(file_name)[1].lower()
                    if ext not in allowed_extensions:
                        results.append({
                            'file': file_name,
                            'status': 'skipped',
                            'message': f'Qo\'llab-quvvatlanmaydigan format: {ext}'
                        })
                        continue
                    
                    try:
                        # Faylni chiqarish
                        zf.extract(file_name, temp_dir)
                        extracted_path = os.path.join(temp_dir, file_name)
                        
                        # Fayl kontentini o'qish
                        content = ""
                        base_name = os.path.basename(file_name)
                        title_from_file = os.path.splitext(base_name)[0]
                        
                        if ext == '.txt':
                            try:
                                with open(extracted_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                            except:
                                with open(extracted_path, 'r', encoding='latin-1') as f:
                                    content = f.read()
                        
                        elif ext == '.pdf':
                            try:
                                from pypdf import PdfReader
                                reader = PdfReader(extracted_path)
                                for page in reader.pages[:10]:  # Birinchi 10 sahifa
                                    content += page.extract_text() or ""
                            except:
                                content = ""
                        
                        elif ext in ['.docx', '.doc']:
                            try:
                                from docx import Document
                                doc = Document(extracted_path)
                                for para in doc.paragraphs[:50]:  # Birinchi 50 paragraf
                                    content += para.text + "\n"
                            except:
                                content = ""
                        
                        # AI bilan ma'lumotlarni aniqlash
                        book_info = ai_extract_book_info(content, base_name)
                        
                        if book_info:
                            author_name = book_info.get('author', '').strip()
                            title = book_info.get('title', title_from_file)
                            year = book_info.get('year')
                            description = book_info.get('description', '')
                        else:
                            author_name = ''
                            title = title_from_file
                            year = None
                            description = ''
                        
                        # Agar muallif nomi bo'sh yoki "noma'lum" bo'lsa
                        if not author_name or author_name.lower() in ['noma\'lum', 'noma\'lum muallif', 'unknown', 'nomalum']:
                            author_name = 'Noma\'lum muallif'
                        
                        # Agar tavsif bo'sh bo'lsa, birinchi 5 qatorni olish
                        if not description and content:
                            lines = content.strip().split('\n')[:5]
                            description = '\n'.join(line.strip() for line in lines if line.strip())
                            if len(description) > 500:
                                description = description[:500] + '...'
                        
                        # Muallif topish yoki yaratish (iexact bilan - katta-kichik harf farq qilmaydi)
                        author = Author.objects.filter(name__iexact=author_name).first()
                        if not author:
                            author = Author.objects.create(name=author_name)
                        
                        # Kitob mavjudligini tekshirish (nom va muallif bo'yicha)
                        existing_book = Book.objects.filter(
                            title__iexact=title,
                            author=author
                        ).first()
                        
                        if existing_book:
                            results.append({
                                'file': base_name,
                                'status': 'exists',
                                'message': f'Kitob bazada mavjud: "{title}" ({author_name})'
                            })
                            continue
                        
                        # Kitobni saqlash
                        from django.core.files.base import ContentFile
                        
                        with open(extracted_path, 'rb') as f:
                            file_content = f.read()
                        
                        book = Book(
                            title=title,
                            description=description if description else '',
                            author=author,
                            year_written=year if year else None
                        )
                        book.file.save(base_name, ContentFile(file_content))
                        book.save()
                        
                        # Sahifalarni saqlash
                        try:
                            book.save_pages_from_file()
                        except:
                            pass
                        
                        results.append({
                            'file': base_name,
                            'status': 'success',
                            'title': title,
                            'author': author_name,
                            'year': year,
                            'description': description[:100] + '...' if len(description) > 100 else description
                        })
                        
                    except Exception as e:
                        results.append({
                            'file': file_name,
                            'status': 'error',
                            'message': str(e)
                        })
            
        return JsonResponse({
            'success': True,
            'total': len(results),
            'results': results
        })
        
    except zipfile.BadZipFile:
        return JsonResponse({'error': 'Noto\'g\'ri ZIP fayl'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# ============== YANGI FUNKSIYALAR ==============

def book_detail(request, book_id):
    """Kitob haqida batafsil sahifa"""
    from .models import Book, BookRating, Favorite, ReadingProgress, SearchQuery
    from django.shortcuts import get_object_or_404
    
    book = get_object_or_404(Book, id=book_id)
    book.views_count += 1
    book.save(update_fields=['views_count'])
    
    ratings = book.ratings.all()[:10]
    is_favorite = False
    reading_progress = None
    
    if request.user.is_authenticated:
        is_favorite = Favorite.objects.filter(user=request.user, book=book).exists()
        reading_progress = ReadingProgress.objects.filter(user=request.user, book=book).first()
    
    # O'xshash kitoblar
    similar_books = Book.objects.filter(
        author=book.author
    ).exclude(id=book.id)[:4]
    
    # Kategoriyalar bo'yicha o'xshash kitoblar (ManyToMany)
    book_categories = book.categories.all()
    if book_categories.exists():
        category_books = Book.objects.filter(
            categories__in=book_categories
        ).exclude(id=book.id).distinct()[:4]
        similar_books = list(similar_books) + list(category_books)
        similar_books = similar_books[:6]
    
    context = {
        'book': book,
        'ratings': ratings,
        'is_favorite': is_favorite,
        'reading_progress': reading_progress,
        'similar_books': similar_books,
        'total_pages': book.pages.count(),
    }
    return render(request, 'blog/book_detail.html', context)


@login_required
def read_book(request, book_id):
    """Kitobni online o'qish - DOCX asl ko'rinishda, PDF LibreOffice bilan"""
    from .models import Book, ReadingProgress
    from django.shortcuts import get_object_or_404
    from django.conf import settings
    import os
    import subprocess
    
    book = get_object_or_404(Book, id=book_id)
    pages = book.pages.all()
    
    current_page = int(request.GET.get('page', 1))
    
    # Fayl mavjud bo'lsa uni ko'rsatish
    file_url = None
    show_file = False
    is_docx = False
    
    if book.file:
        file_path = book.file.path
        file_name = book.file.name.lower()
        
        if file_name.endswith('.pdf'):
            # PDF to'g'ridan-to'g'ri ko'rsatiladi
            file_url = book.file.url
            show_file = True
            is_docx = False
        elif file_name.endswith(('.docx', '.doc')):
            # DOCX/DOC to'g'ridan-to'g'ri docx-preview.js orqali ko'rsatiladi
            file_url = book.file.url
            show_file = True
            is_docx = True
        elif file_name.endswith(('.odt', '.txt', '.rtf')):
            # Boshqa formatlar - faqat matn rejimida ko'rsatish
            show_file = False
            is_docx = False
    
    if request.user.is_authenticated:
        progress, created = ReadingProgress.objects.get_or_create(
            user=request.user,
            book=book,
            defaults={'total_pages': pages.count()}
        )
        if not created and current_page > progress.current_page:
            progress.current_page = current_page
            progress.total_pages = pages.count()
            progress.save()
    
    context = {
        'book': book,
        'pages': pages,
        'current_page': current_page,
        'total_pages': pages.count(),
        'file_url': file_url,
        'show_file': show_file,
        'is_docx': is_docx,
    }
    return render(request, 'blog/read_book.html', context)


@csrf_exempt
def rate_book(request):
    """Kitobga baho qo'yish"""
    from .models import Book, BookRating
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST'}, status=400)
    
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        rating = data.get('rating')
        comment = data.get('comment', '')
        name = data.get('name', '')
        
        book = Book.objects.get(id=book_id)
        
        BookRating.objects.create(
            book=book,
            user=request.user if request.user.is_authenticated else None,
            name=name,
            rating=rating,
            comment=comment
        )
        
        return JsonResponse({
            'success': True,
            'average_rating': book.average_rating,
            'total_ratings': book.total_ratings
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
def toggle_favorite(request):
    """Sevimlilar ro'yxatiga qo'shish/o'chirish"""
    from .models import Book, Favorite
    import json
    
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Tizimga kiring'}, status=401)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST'}, status=400)
    
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        book = Book.objects.get(id=book_id)
        
        favorite, created = Favorite.objects.get_or_create(
            user=request.user,
            book=book
        )
        
        if not created:
            favorite.delete()
            return JsonResponse({'success': True, 'is_favorite': False})
        
        return JsonResponse({'success': True, 'is_favorite': True})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
def save_reading_progress(request):
    """O'qish progressini saqlash"""
    from .models import Book, ReadingProgress
    import json
    
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Tizimga kiring'}, status=401)
    
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        current_page = data.get('current_page')
        
        book = Book.objects.get(id=book_id)
        total_pages = book.pages.count()
        
        progress, _ = ReadingProgress.objects.update_or_create(
            user=request.user,
            book=book,
            defaults={
                'current_page': current_page,
                'total_pages': total_pages,
                'is_completed': current_page >= total_pages
            }
        )
        
        return JsonResponse({
            'success': True,
            'progress_percent': progress.progress_percent
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
def get_book_summary(request):
    """AI bilan kitob xulosasini olish"""
    from .models import Book, BookSummary
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST'}, status=400)
    
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        book = Book.objects.get(id=book_id)
        
        # Mavjud xulosa bormi?
        try:
            summary = book.summary
            return JsonResponse({
                'success': True,
                'summary': summary.short_summary,
                'key_points': summary.key_points
            })
        except BookSummary.DoesNotExist:
            pass
        
        # AI bilan xulosa yaratish
        if not groq_client:
            return JsonResponse({'error': 'AI xizmati mavjud emas'}, status=400)
        
        content = book.content[:5000] if book.content else ""
        if not content and book.pages.exists():
            content = "\n".join([p.text for p in book.pages.all()[:5]])[:5000]
        
        if not content:
            return JsonResponse({'error': 'Kitob matni mavjud emas'}, status=400)
        
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": "Kitob xulosasini o'zbek tilida yoz. Qisqa va aniq bo'l."
                },
                {
                    "role": "user",
                    "content": f"Kitob: {book.title}\nMuallif: {book.author.name}\n\nMatn:\n{content}\n\nQisqa xulosa yoz (3-5 jumla)."
                }
            ],
            temperature=0.5,
            max_tokens=500
        )
        
        summary_text = response.choices[0].message.content
        
        BookSummary.objects.create(
            book=book,
            short_summary=summary_text
        )
        
        return JsonResponse({
            'success': True,
            'summary': summary_text
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
def ask_about_book(request):
    """Kitob haqida savol berish"""
    from .models import Book
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST'}, status=400)
    
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        question = data.get('question', '').strip()
        
        if not question:
            return JsonResponse({'error': 'Savol bo\'sh'}, status=400)
        
        book = Book.objects.get(id=book_id)
        
        if not groq_client:
            return JsonResponse({'error': 'AI xizmati mavjud emas'}, status=400)
        
        content = book.content[:6000] if book.content else ""
        if not content and book.pages.exists():
            content = "\n".join([p.text for p in book.pages.all()[:8]])[:6000]
        
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": f"Sen {book.title} kitobi bo'yicha mutaxassissan. Savolga kitob asosida javob ber. O'zbek tilida qisqa va aniq javob ber."
                },
                {
                    "role": "user",
                    "content": f"Kitob matni:\n{content}\n\nSavol: {question}"
                }
            ],
            temperature=0.7,
            max_tokens=600
        )
        
        answer = response.choices[0].message.content
        
        return JsonResponse({
            'success': True,
            'answer': answer
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
def get_similar_books(request):
    """O'xshash kitoblarni topish"""
    from .models import Book
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST'}, status=400)
    
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        book = Book.objects.get(id=book_id)
        
        # Muallif va kategoriya bo'yicha o'xshash (ManyToMany)
        similar = Book.objects.filter(
            models.Q(author=book.author) | models.Q(categories__in=book.categories.all())
        ).exclude(id=book.id).distinct()[:8]
        
        results = [{
            'id': b.id,
            'title': b.title,
            'author': b.author.name,
            'cover': b.cover_image.url if b.cover_image else None,
            'rating': b.average_rating
        } for b in similar]
        
        return JsonResponse({'success': True, 'books': results})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


def categories_list(request):
    """Kategoriyalar ro'yxati"""
    from .models import Category
    categories = Category.objects.all()
    return render(request, 'blog/categories.html', {'categories': categories})


def category_books(request, slug):
    """Kategoriya bo'yicha kitoblar - all_books ga yo'naltirish"""
    from django.shortcuts import redirect
    return redirect(f'/kitoblar/?category={slug}')


# ESKIRGAN: top_books, new_books, favorites_list - bular endi 728-qatordagi yangi versiyada
# Eski kodni saqlab qolish (lekin ishlatilmaydi)


@login_required
def delete_account(request):
    """Foydalanuvchi akkauntini va uning fayllarini o'chirish"""
    from .models import Image, File
    import os
    
    if request.method == 'POST':
        user = request.user
        
        # Foydalanuvchi yuklagan rasmlarni o'chirish
        user_images = Image.objects.filter(user=user)
        for img in user_images:
            if img.image and os.path.exists(img.image.path):
                try:
                    os.remove(img.image.path)
                except:
                    pass
        user_images.delete()
        
        # Foydalanuvchi yuklagan fayllarni o'chirish
        user_files = File.objects.filter(user=user)
        for f in user_files:
            if f.file and os.path.exists(f.file.path):
                try:
                    os.remove(f.file.path)
                except:
                    pass
        user_files.delete()
        
        # Akkauntni o'chirish
        from django.contrib.auth import logout
        logout(request)
        user.delete()
        
        return redirect('boshlash')
    
    return redirect('profile')


def user_profile(request):
    """Foydalanuvchi profili - user_profile_view ga yo'naltirish"""
    if not request.user.is_authenticated:
        from django.shortcuts import redirect
        return redirect('login')
    
    # O'z profilini ko'rish uchun social profile ga yo'naltirish
    return redirect('user_profile', username=request.user.username)


def offline_page(request):
    """Offline sahifa (PWA uchun)"""
    return render(request, 'blog/offline.html')


def privacy_policy(request):
    """Maxfiylik siyosati sahifasi"""
    return render(request, 'blog/privacy_policy.html')


def terms_of_service(request):
    """Foydalanish shartlari sahifasi"""
    return render(request, 'blog/terms_of_service.html')


def about(request):
    """Biz haqimizda sahifasi"""
    return render(request, 'blog/about.html')


# ===== PWA (Progressive Web App) =====
from django.http import HttpResponse
import json

def pwa_manifest(request):
    """PWA manifest.json - root'dan xizmat qilish"""
    manifest = {
        "name": "Cloudstore - Kutubxona",
        "short_name": "Cloudstore",
        "description": "Bepul elektron kitoblar kutubxonasi",
        "start_url": "/",
        "display": "standalone",
        "background_color": "#667eea",
        "theme_color": "#667eea",
        "orientation": "portrait-primary",
        "scope": "/",
        "lang": "uz",
        "icons": [
            {"src": "/static/icons/icon-72x72.png", "sizes": "72x72", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-96x96.png", "sizes": "96x96", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-128x128.png", "sizes": "128x128", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-144x144.png", "sizes": "144x144", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-152x152.png", "sizes": "152x152", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-192x192.png", "sizes": "192x192", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-384x384.png", "sizes": "384x384", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icons/icon-512x512.png", "sizes": "512x512", "type": "image/png", "purpose": "any maskable"}
        ]
    }
    return HttpResponse(json.dumps(manifest), content_type='application/manifest+json')


def pwa_service_worker(request):
    """PWA Service Worker - root'dan xizmat qilish"""
    sw_content = '''
const CACHE_NAME = 'cloudstore-v2';
const OFFLINE_URL = '/offline/';

const STATIC_ASSETS = [
    '/',
    '/offline/',
    '/static/blog/style.css'
];

self.addEventListener('install', event => {
    event.waitUntil(
        caches.open(CACHE_NAME)
            .then(cache => cache.addAll(STATIC_ASSETS))
            .then(() => self.skipWaiting())
    );
});

self.addEventListener('activate', event => {
    event.waitUntil(
        caches.keys().then(cacheNames => {
            return Promise.all(
                cacheNames
                    .filter(name => name !== CACHE_NAME)
                    .map(name => caches.delete(name))
            );
        }).then(() => self.clients.claim())
    );
});

self.addEventListener('fetch', event => {
    if (event.request.mode === 'navigate') {
        event.respondWith(
            fetch(event.request)
                .catch(() => caches.match(OFFLINE_URL))
        );
    }
});
'''
    return HttpResponse(sw_content.strip(), content_type='application/javascript')


def pwa_offline(request):
    """Offline sahifa"""
    return render(request, 'blog/offline.html')


# ===== BARCODE SCANNER =====

def barcode_scanner(request):
    """Barcode skanerlash sahifasi"""
    from .models import ProductScanHistory
    
    recent_scans = []
    if request.user.is_authenticated:
        recent_scans = ProductScanHistory.objects.filter(user=request.user)[:10]
    
    return render(request, 'blog/barcode_scanner.html', {
        'recent_scans': recent_scans
    })


@csrf_exempt
def barcode_lookup(request):
    """Barcode orqali mahsulot ma'lumotlarini olish API"""
    from .models import Product, ProductScanHistory
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST so\'rov qabul qilinadi'}, status=400)
    
    try:
        data = json.loads(request.body)
        barcode = data.get('barcode', '').strip()
        
        if not barcode:
            return JsonResponse({'error': 'Barcode kiritilmagan'}, status=400)
        
        # Mahsulotni qidirish
        try:
            product = Product.objects.get(barcode=barcode)
            
            # Skanerlash sonini oshirish
            product.scan_count += 1
            product.save(update_fields=['scan_count'])
            
            # Tarixga qo'shish
            if request.user.is_authenticated:
                ProductScanHistory.objects.create(
                    user=request.user,
                    product=product
                )
            
            return JsonResponse({
                'success': True,
                'found': True,
                'product': {
                    'id': product.id,
                    'barcode': product.barcode,
                    'name': product.name,
                    'description': product.description,
                    'price': str(product.price) if product.price else None,
                    'currency': product.currency,
                    'category': product.category,
                    'brand': product.brand,
                    'manufacturer': product.manufacturer,
                    'country': product.country,
                    'weight': product.weight,
                    'image': product.image.url if product.image else None,
                    'image_url': product.image_url,
                    'ingredients': product.ingredients,
                    'expiry_info': product.expiry_info,
                    'scan_count': product.scan_count,
                    
                    # Ozuqaviy qiymatlar
                    'calories': product.calories,
                    'fat': product.fat,
                    'saturated_fat': product.saturated_fat,
                    'carbohydrates': product.carbohydrates,
                    'sugars': product.sugars,
                    'fiber': product.fiber,
                    'proteins': product.proteins,
                    'salt': product.salt,
                    'sodium': product.sodium,
                    
                    # Vitaminlar
                    'vitamin_a': product.vitamin_a,
                    'vitamin_c': product.vitamin_c,
                    'vitamin_d': product.vitamin_d,
                    'vitamin_e': product.vitamin_e,
                    'vitamin_b1': product.vitamin_b1,
                    'vitamin_b2': product.vitamin_b2,
                    'vitamin_b6': product.vitamin_b6,
                    'vitamin_b12': product.vitamin_b12,
                    
                    # Minerallar
                    'calcium': product.calcium,
                    'iron': product.iron,
                    'magnesium': product.magnesium,
                    'zinc': product.zinc,
                    'potassium': product.potassium,
                    
                    # Sog'liq ko'rsatkichlari
                    'nutriscore_grade': product.nutriscore_grade,
                    'nova_group': product.nova_group,
                    'ecoscore_grade': product.ecoscore_grade,
                    
                    # Allergenlar va ogohlantirishlar
                    'allergens': product.allergens,
                    'additives': product.additives,
                    'warnings': product.warnings,
                }
            })
        except Product.DoesNotExist:
            # Tashqi API dan qidirish (Open Food Facts)
            external_data = fetch_from_external_api(barcode)
            
            if external_data:
                # Yangi mahsulot yaratish to'liq ma'lumotlar bilan
                product = Product.objects.create(
                    barcode=barcode,
                    name=external_data.get('name', f'Mahsulot {barcode}'),
                    description=external_data.get('description', ''),
                    brand=external_data.get('brand', ''),
                    category=external_data.get('category', ''),
                    country=external_data.get('country', ''),
                    ingredients=external_data.get('ingredients', ''),
                    weight=external_data.get('weight', ''),
                    image_url=external_data.get('image_url', ''),
                    
                    # Ozuqaviy qiymatlar
                    calories=external_data.get('calories'),
                    fat=external_data.get('fat'),
                    saturated_fat=external_data.get('saturated_fat'),
                    carbohydrates=external_data.get('carbohydrates'),
                    sugars=external_data.get('sugars'),
                    fiber=external_data.get('fiber'),
                    proteins=external_data.get('proteins'),
                    salt=external_data.get('salt'),
                    sodium=external_data.get('sodium'),
                    
                    # Vitaminlar
                    vitamin_a=external_data.get('vitamin_a'),
                    vitamin_c=external_data.get('vitamin_c'),
                    vitamin_d=external_data.get('vitamin_d'),
                    vitamin_e=external_data.get('vitamin_e'),
                    vitamin_b1=external_data.get('vitamin_b1'),
                    vitamin_b2=external_data.get('vitamin_b2'),
                    vitamin_b6=external_data.get('vitamin_b6'),
                    vitamin_b12=external_data.get('vitamin_b12'),
                    
                    # Minerallar
                    calcium=external_data.get('calcium'),
                    iron=external_data.get('iron'),
                    magnesium=external_data.get('magnesium'),
                    zinc=external_data.get('zinc'),
                    potassium=external_data.get('potassium'),
                    
                    # Sog'liq ko'rsatkichlari
                    nutriscore_grade=external_data.get('nutriscore_grade', ''),
                    nova_group=external_data.get('nova_group'),
                    ecoscore_grade=external_data.get('ecoscore_grade', ''),
                    
                    # Allergenlar
                    allergens=external_data.get('allergens', ''),
                    additives=external_data.get('additives', ''),
                    warnings=external_data.get('warnings', ''),
                )
                
                if request.user.is_authenticated:
                    ProductScanHistory.objects.create(user=request.user, product=product)
                
                return JsonResponse({
                    'success': True,
                    'found': True,
                    'source': 'external',
                    'product': {
                        'id': product.id,
                        'barcode': product.barcode,
                        'name': product.name,
                        'description': product.description,
                        'brand': product.brand,
                        'category': product.category,
                        'country': product.country,
                        'ingredients': product.ingredients,
                        'weight': product.weight,
                        'image_url': product.image_url,
                        
                        # Ozuqaviy qiymatlar
                        'calories': product.calories,
                        'fat': product.fat,
                        'saturated_fat': product.saturated_fat,
                        'carbohydrates': product.carbohydrates,
                        'sugars': product.sugars,
                        'fiber': product.fiber,
                        'proteins': product.proteins,
                        'salt': product.salt,
                        
                        # Vitaminlar
                        'vitamin_a': product.vitamin_a,
                        'vitamin_c': product.vitamin_c,
                        'vitamin_d': product.vitamin_d,
                        
                        # Sog'liq ko'rsatkichlari
                        'nutriscore_grade': product.nutriscore_grade,
                        'nova_group': product.nova_group,
                        
                        # Allergenlar va ogohlantirishlar
                        'allergens': product.allergens,
                        'additives': product.additives,
                        'warnings': product.warnings,
                    }
                })
            
            return JsonResponse({
                'success': True,
                'found': False,
                'barcode': barcode,
                'message': 'Mahsulot topilmadi. Siz uni qo\'shishingiz mumkin.'
            })
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Noto\'g\'ri JSON format'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def fetch_from_external_api(barcode):
    """Open Food Facts API dan mahsulot ma'lumotlarini olish"""
    import requests
    
    try:
        # Open Food Facts API
        url = f"https://world.openfoodfacts.org/api/v0/product/{barcode}.json"
        response = requests.get(url, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('status') == 1:
                product = data.get('product', {})
                nutriments = product.get('nutriments', {})
                
                # Allergenlarni olish
                allergens_tags = product.get('allergens_tags', [])
                allergens = ', '.join([a.replace('en:', '').replace('-', ' ').title() for a in allergens_tags])
                
                # Qo'shimchalarni olish (E raqamlar)
                additives_tags = product.get('additives_tags', [])
                additives = ', '.join([a.replace('en:', '').upper() for a in additives_tags])
                
                # Ogohlantirishlar
                warnings_list = []
                if nutriments.get('sugars_100g', 0) and nutriments.get('sugars_100g', 0) > 22.5:
                    warnings_list.append("⚠️ Shakar miqdori yuqori")
                if nutriments.get('salt_100g', 0) and nutriments.get('salt_100g', 0) > 1.5:
                    warnings_list.append("⚠️ Tuz miqdori yuqori")
                if nutriments.get('saturated-fat_100g', 0) and nutriments.get('saturated-fat_100g', 0) > 5:
                    warnings_list.append("⚠️ To'yingan yog' miqdori yuqori")
                nova = product.get('nova_group')
                if nova and nova == 4:
                    warnings_list.append("⚠️ Yuqori darajada qayta ishlangan mahsulot")
                
                return {
                    'name': product.get('product_name', '') or product.get('product_name_en', ''),
                    'description': product.get('generic_name', ''),
                    'brand': product.get('brands', ''),
                    'category': product.get('categories', ''),
                    'country': product.get('countries', ''),
                    'ingredients': product.get('ingredients_text', ''),
                    'weight': product.get('quantity', ''),
                    'image_url': product.get('image_url', ''),
                    
                    # Asosiy ozuqaviy qiymatlar
                    'calories': nutriments.get('energy-kcal_100g'),
                    'fat': nutriments.get('fat_100g'),
                    'saturated_fat': nutriments.get('saturated-fat_100g'),
                    'carbohydrates': nutriments.get('carbohydrates_100g'),
                    'sugars': nutriments.get('sugars_100g'),
                    'fiber': nutriments.get('fiber_100g'),
                    'proteins': nutriments.get('proteins_100g'),
                    'salt': nutriments.get('salt_100g'),
                    'sodium': nutriments.get('sodium_100g'),
                    
                    # Vitaminlar
                    'vitamin_a': nutriments.get('vitamin-a_100g'),
                    'vitamin_c': nutriments.get('vitamin-c_100g'),
                    'vitamin_d': nutriments.get('vitamin-d_100g'),
                    'vitamin_e': nutriments.get('vitamin-e_100g'),
                    'vitamin_b1': nutriments.get('vitamin-b1_100g'),
                    'vitamin_b2': nutriments.get('vitamin-b2_100g'),
                    'vitamin_b6': nutriments.get('vitamin-b6_100g'),
                    'vitamin_b12': nutriments.get('vitamin-b12_100g'),
                    
                    # Minerallar
                    'calcium': nutriments.get('calcium_100g'),
                    'iron': nutriments.get('iron_100g'),
                    'magnesium': nutriments.get('magnesium_100g'),
                    'zinc': nutriments.get('zinc_100g'),
                    'potassium': nutriments.get('potassium_100g'),
                    
                    # Sog'liq baholari
                    'nutriscore_grade': product.get('nutriscore_grade', '').upper(),
                    'nova_group': product.get('nova_group'),
                    'ecoscore_grade': product.get('ecoscore_grade', '').upper(),
                    
                    # Allergenlar va ogohlantirishlar
                    'allergens': allergens,
                    'additives': additives,
                    'warnings': '\n'.join(warnings_list),
                }
    except Exception:
        pass
    
    return None


@csrf_exempt  
def barcode_add_product(request):
    """Yangi mahsulot qo'shish"""
    from .models import Product
    import json
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Faqat POST'}, status=400)
    
    try:
        data = json.loads(request.body)
        barcode = data.get('barcode', '').strip()
        name = data.get('name', '').strip()
        
        if not barcode or not name:
            return JsonResponse({'error': 'Barcode va nom majburiy'}, status=400)
        
        # Mavjudligini tekshirish
        if Product.objects.filter(barcode=barcode).exists():
            return JsonResponse({'error': 'Bu barcode allaqachon mavjud'}, status=400)
        
        product = Product.objects.create(
            barcode=barcode,
            name=name,
            description=data.get('description', ''),
            price=data.get('price') if data.get('price') else None,
            currency=data.get('currency', 'UZS'),
            category=data.get('category', ''),
            brand=data.get('brand', ''),
            manufacturer=data.get('manufacturer', ''),
            country=data.get('country', ''),
            weight=data.get('weight', ''),
            ingredients=data.get('ingredients', ''),
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Mahsulot muvaffaqiyatli qo\'shildi',
            'product_id': product.id
        })
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Noto\'g\'ri JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def barcode_history(request):
    """Skanerlash tarixi"""
    from .models import ProductScanHistory
    
    if not request.user.is_authenticated:
        return render(request, 'blog/barcode_history.html', {'scans': []})
    
    scans = ProductScanHistory.objects.filter(user=request.user).select_related('product')[:50]
    return render(request, 'blog/barcode_history.html', {'scans': scans})
