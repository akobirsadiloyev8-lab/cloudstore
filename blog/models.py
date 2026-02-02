from django.db import models
from django.contrib.auth.models import User
from django.db.models import Avg


class Category(models.Model):
    """Kitob janrlari/kategoriyalari"""
    name = models.CharField(max_length=100, verbose_name="Kategoriya nomi")
    slug = models.SlugField(unique=True)
    icon = models.CharField(max_length=50, default='fa-book', verbose_name="Ikonka (FontAwesome)")
    description = models.TextField(blank=True, verbose_name="Tavsif")
    
    class Meta:
        verbose_name = "Kategoriya"
        verbose_name_plural = "Kategoriyalar"
        ordering = ['name']
    
    def __str__(self):
        return self.name


class BookPage(models.Model):
    book = models.ForeignKey('Book', on_delete=models.CASCADE, related_name='pages', verbose_name="Kitob")
    page_number = models.PositiveIntegerField(verbose_name="Sahifa raqami")
    text = models.TextField(verbose_name="Sahifa matni")

    class Meta:
        unique_together = ('book', 'page_number')
        ordering = ['page_number']
        verbose_name = "Kitob sahifasi"
        verbose_name_plural = "Kitob sahifalari"

    def __str__(self):
        return f"{self.book.title} - {self.page_number}-sahifa"

class Image(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    title = models.CharField(max_length=100, blank=True)
    image = models.ImageField(upload_to='images/')
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.title or self.image.name

class File(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    file = models.FileField(upload_to='files/')
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.file.name

class Feedback(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    name = models.CharField(max_length=100, blank=True, verbose_name="Ism")
    message = models.TextField(verbose_name="Xabar")
    is_read = models.BooleanField(default=False, verbose_name="O'qildi")
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Yuborilgan vaqt")

    class Meta:
        verbose_name = "Fikr-mulohaza"
        verbose_name_plural = "Fikr-mulohazalar"
        ordering = ['-created_at']

    def __str__(self):
        return f"{self.name or 'Anonim'}: {self.message[:50]}"


class Author(models.Model):
    """Adib modeli"""
    name = models.CharField(max_length=200, verbose_name="Adib ismi")
    bio = models.TextField(blank=True, verbose_name="Tarjimai hol")
    birth_year = models.IntegerField(null=True, blank=True, verbose_name="Tug'ilgan yili")
    death_year = models.IntegerField(null=True, blank=True, verbose_name="Vafot etgan yili")
    image = models.ImageField(upload_to='authors/', blank=True, null=True, verbose_name="Rasm")
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = "Adib"
        verbose_name_plural = "Adiblar"
        ordering = ['name']

    def __str__(self):
        return self.name



class Book(models.Model):
    """Kitob modeli"""
    author = models.ForeignKey(Author, on_delete=models.CASCADE, related_name='books', verbose_name="Adib")
    title = models.CharField(max_length=300, verbose_name="Kitob nomi")
    description = models.TextField(blank=True, verbose_name="Tavsif")
    year_written = models.IntegerField(null=True, blank=True, verbose_name="Yozilgan yili")
    file = models.FileField(upload_to='books/', blank=True, null=True, verbose_name="Fayl (PDF/DOCX)")
    content = models.TextField(blank=True, verbose_name="Kitob matni", help_text="Fayldan avtomatik o'qilgan to'liq matn")
    cover_image = models.ImageField(upload_to='book_covers/', blank=True, null=True, verbose_name="Muqova rasmi")
    categories = models.ManyToManyField(Category, blank=True, related_name='books', verbose_name="Kategoriyalar")
    views_count = models.PositiveIntegerField(default=0, verbose_name="Ko'rishlar soni")
    created_at = models.DateTimeField(auto_now_add=True)
    
    @property
    def average_rating(self):
        """O'rtacha reyting"""
        avg = self.ratings.aggregate(Avg('rating'))['rating__avg']
        return round(avg, 1) if avg else 0
    
    @property
    def total_ratings(self):
        """Baholar soni"""
        return self.ratings.count()

    def save_pages_from_file(self):
        """Fayldan matnni sahifalarga bo‘lib BookPage obyektlari sifatida saqlash"""
        from .models import BookPage
        import os
        import subprocess
        import tempfile
        
        if not self.file:
            return
        file_path = self.file.path
        file_name = self.file.name.lower()
        # Avval eski sahifalarni o'chir
        self.pages.all().delete()
        
        if file_name.endswith('.pdf'):
            all_text = ""
            pages_created = False
            
            # 1-usul: pdfplumber - eng yaxshi strukturani saqlaydi
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        # Asosiy matnni olish
                        text = page.extract_text() or ''
                        
                        # Jadvallarni ham olish (agar bo'lsa)
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                for row in table:
                                    if row:
                                        row_text = ' | '.join([str(cell) if cell else '' for cell in row])
                                        text += '\n' + row_text
                        
                        all_text += text
                        if text.strip():
                            BookPage.objects.create(book=self, page_number=i+1, text=text)
                            pages_created = True
                print(f"pdfplumber bilan {len(pdf.pages)} sahifa o'qildi")
            except Exception as e:
                print(f"pdfplumber xatosi: {e}")
            
            # 2-usul: pypdf - agar pdfplumber ishlamasa
            if not all_text.strip() and not pages_created:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text() or ''
                        all_text += text
                        if text.strip():
                            BookPage.objects.create(book=self, page_number=i+1, text=text)
                            pages_created = True
                    print(f"pypdf bilan {len(reader.pages)} sahifa o'qildi")
                except Exception as e:
                    print(f"pypdf xatosi: {e}")
            
            # 3-usul: LibreOffice - oxirgi variant (skaner qilingan PDF uchun)
            if not all_text.strip() and not pages_created:
                try:
                    with tempfile.TemporaryDirectory() as temp_dir:
                        result = subprocess.run([
                            'soffice', '--headless', '--convert-to', 'txt:Text',
                            '--outdir', temp_dir, file_path
                        ], capture_output=True, timeout=120)
                        
                        base_name = os.path.splitext(os.path.basename(file_path))[0]
                        txt_path = os.path.join(temp_dir, f'{base_name}.txt')
                        
                        if os.path.exists(txt_path):
                            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                            
                            if content.strip():
                                self.pages.all().delete()
                                page_size = 2000
                                pages_text = [content[i:i+page_size] for i in range(0, len(content), page_size)]
                                
                                for i, page_text in enumerate(pages_text):
                                    if page_text.strip():
                                        BookPage.objects.create(book=self, page_number=i+1, text=page_text)
                                
                                print(f"LibreOffice bilan {len(pages_text)} sahifa olindi")
                except Exception as e:
                    print(f"LibreOffice PDF->TXT xatosi: {e}")
                    
        elif file_name.endswith('.docx'):
            from docx import Document
            doc = Document(file_path)
            # Har 40 paragraf = 1 sahifa
            page_size = 40
            paragraphs = [p.text for p in doc.paragraphs]
            for i in range(0, len(paragraphs), page_size):
                page_text = '\n'.join(paragraphs[i:i+page_size])
                BookPage.objects.create(book=self, page_number=(i//page_size)+1, text=page_text)
        elif file_name.endswith('.doc'):
            # .doc fayllar uchun LibreOffice
            try:
                with tempfile.TemporaryDirectory() as temp_dir:
                    result = subprocess.run([
                        'soffice', '--headless', '--convert-to', 'txt:Text',
                        '--outdir', temp_dir, file_path
                    ], capture_output=True, timeout=120)
                    
                    base_name = os.path.splitext(os.path.basename(file_path))[0]
                    txt_path = os.path.join(temp_dir, f'{base_name}.txt')
                    
                    if os.path.exists(txt_path):
                        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        
                        page_size = 2000
                        pages_text = [content[i:i+page_size] for i in range(0, len(content), page_size)]
                        
                        for i, page_text in enumerate(pages_text):
                            if page_text.strip():
                                BookPage.objects.create(book=self, page_number=i+1, text=page_text)
            except Exception as e:
                print(f"LibreOffice DOC->TXT xatosi: {e}")
        elif file_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            # Har 50 satr = 1 sahifa
            page_size = 50
            for i in range(0, len(lines), page_size):
                page_text = ''.join(lines[i:i+page_size])
                BookPage.objects.create(book=self, page_number=(i//page_size)+1, text=page_text)
        elif file_name.endswith(('.odt', '.rtf', '.ppt', '.pptx', '.xls', '.xlsx')):
            # Boshqa formatlar uchun LibreOffice
            try:
                with tempfile.TemporaryDirectory() as temp_dir:
                    result = subprocess.run([
                        'soffice', '--headless', '--convert-to', 'txt:Text',
                        '--outdir', temp_dir, file_path
                    ], capture_output=True, timeout=120)
                    
                    base_name = os.path.splitext(os.path.basename(file_path))[0]
                    txt_path = os.path.join(temp_dir, f'{base_name}.txt')
                    
                    if os.path.exists(txt_path):
                        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        
                        page_size = 2000
                        pages_text = [content[i:i+page_size] for i in range(0, len(content), page_size)]
                        
                        for i, page_text in enumerate(pages_text):
                            if page_text.strip():
                                BookPage.objects.create(book=self, page_number=i+1, text=page_text)
            except Exception as e:
                print(f"LibreOffice convert xatosi: {e}")

    class Meta:
        verbose_name = "Kitob"
        verbose_name_plural = "Kitoblar"
        ordering = ['title']

    def __str__(self):
        return f"{self.title} - {self.author.name}"

    def extract_text_from_file(self):
        """Fayldan matnni o'qib olish"""
        if not self.file:
            return ""
        
        file_path = self.file.path
        file_name = self.file.name.lower()
        text = ""
        
        try:
            if file_name.endswith('.pdf'):
                text = self._extract_from_pdf(file_path)
            elif file_name.endswith('.docx'):
                text = self._extract_from_docx(file_path)
            elif file_name.endswith('.txt'):
                text = self._extract_from_txt(file_path)
        except Exception as e:
            text = f"Faylni o'qishda xatolik: {str(e)}"
        
        return text
    
    def _extract_from_pdf(self, file_path):
        """PDF fayldan matn o'qish"""
        from pypdf import PdfReader
        text = ""
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    
    def _extract_from_docx(self, file_path):
        """DOCX fayldan matn o'qish"""
        from docx import Document
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_txt(self, file_path):
        """TXT fayldan matn o'qish"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()


class BookRating(models.Model):
    """Kitobga baho va sharh"""
    book = models.ForeignKey(Book, on_delete=models.CASCADE, related_name='ratings')
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    name = models.CharField(max_length=100, blank=True, verbose_name="Ism")
    rating = models.PositiveSmallIntegerField(choices=[(i, i) for i in range(1, 6)], verbose_name="Baho (1-5)")
    comment = models.TextField(blank=True, verbose_name="Sharh")
    created_at = models.DateTimeField(auto_now_add=True)
    
    class Meta:
        verbose_name = "Kitob bahosi"
        verbose_name_plural = "Kitob baholari"
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.book.title} - {self.rating}⭐ ({self.name or 'Anonim'})"


class Favorite(models.Model):
    """Sevimli kitoblar"""
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='favorites')
    book = models.ForeignKey(Book, on_delete=models.CASCADE, related_name='favorited_by')
    created_at = models.DateTimeField(auto_now_add=True)
    
    class Meta:
        unique_together = ('user', 'book')
        verbose_name = "Sevimli"
        verbose_name_plural = "Sevimlilar"
    
    def __str__(self):
        return f"{self.user.username} - {self.book.title}"


class ReadingProgress(models.Model):
    """O'qish progressi"""
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='reading_progress')
    book = models.ForeignKey(Book, on_delete=models.CASCADE, related_name='readers')
    current_page = models.PositiveIntegerField(default=1, verbose_name="Hozirgi sahifa")
    total_pages = models.PositiveIntegerField(default=1, verbose_name="Jami sahifalar")
    last_read = models.DateTimeField(auto_now=True)
    started_at = models.DateTimeField(auto_now_add=True)
    is_completed = models.BooleanField(default=False, verbose_name="Tugallandi")
    
    class Meta:
        unique_together = ('user', 'book')
        verbose_name = "O'qish progressi"
        verbose_name_plural = "O'qish progresslari"
    
    @property
    def progress_percent(self):
        if self.total_pages > 0:
            return round((self.current_page / self.total_pages) * 100)
        return 0
    
    def __str__(self):
        return f"{self.user.username} - {self.book.title} ({self.progress_percent}%)"


class SearchQuery(models.Model):
    """Qidiruv so'rovlari statistikasi"""
    query = models.CharField(max_length=255, verbose_name="Qidiruv so'zi")
    count = models.PositiveIntegerField(default=1, verbose_name="Qidirilgan soni")
    last_searched = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = "Qidiruv so'rovi"
        verbose_name_plural = "Qidiruv so'rovlari"
        ordering = ['-count']
    
    def __str__(self):
        return f"{self.query} ({self.count})"


class BookSummary(models.Model):
    """Kitob xulosasi (AI tomonidan yaratilgan)"""
    book = models.OneToOneField(Book, on_delete=models.CASCADE, related_name='summary')
    short_summary = models.TextField(verbose_name="Qisqa xulosa")
    key_points = models.TextField(blank=True, verbose_name="Asosiy fikrlar")
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = "Kitob xulosasi"
        verbose_name_plural = "Kitob xulosalari"
    
    def __str__(self):
        return f"Xulosa: {self.book.title}"


class Product(models.Model):
    """Mahsulot modeli - shtrix kod bilan"""
    barcode = models.CharField(max_length=50, unique=True, verbose_name="Shtrix kod (Barcode)")
    name = models.CharField(max_length=300, verbose_name="Mahsulot nomi")
    description = models.TextField(blank=True, verbose_name="Tavsif")
    price = models.DecimalField(max_digits=12, decimal_places=2, null=True, blank=True, verbose_name="Narxi")
    currency = models.CharField(max_length=10, default="UZS", verbose_name="Valyuta")
    category = models.CharField(max_length=100, blank=True, verbose_name="Kategoriya")
    brand = models.CharField(max_length=100, blank=True, verbose_name="Brend")
    manufacturer = models.CharField(max_length=200, blank=True, verbose_name="Ishlab chiqaruvchi")
    country = models.CharField(max_length=100, blank=True, verbose_name="Ishlab chiqarilgan mamlakat")
    weight = models.CharField(max_length=50, blank=True, verbose_name="Og'irligi")
    image = models.ImageField(upload_to='products/', blank=True, null=True, verbose_name="Rasm")
    image_url = models.URLField(max_length=500, blank=True, verbose_name="Rasm URL")
    ingredients = models.TextField(blank=True, verbose_name="Tarkibi")
    nutrition_info = models.TextField(blank=True, verbose_name="Ozuqaviy qiymati")
    expiry_info = models.CharField(max_length=100, blank=True, verbose_name="Yaroqlilik muddati")
    
    # Ozuqaviy qiymatlar (100g uchun)
    calories = models.FloatField(null=True, blank=True, verbose_name="Kaloriya (kcal)")
    fat = models.FloatField(null=True, blank=True, verbose_name="Yog' (g)")
    saturated_fat = models.FloatField(null=True, blank=True, verbose_name="To'yingan yog' (g)")
    carbohydrates = models.FloatField(null=True, blank=True, verbose_name="Uglevodlar (g)")
    sugars = models.FloatField(null=True, blank=True, verbose_name="Shakar (g)")
    fiber = models.FloatField(null=True, blank=True, verbose_name="Kletchatka (g)")
    proteins = models.FloatField(null=True, blank=True, verbose_name="Oqsillar (g)")
    salt = models.FloatField(null=True, blank=True, verbose_name="Tuz (g)")
    sodium = models.FloatField(null=True, blank=True, verbose_name="Natriy (mg)")
    
    # Vitaminlar va minerallar
    vitamin_a = models.FloatField(null=True, blank=True, verbose_name="Vitamin A (µg)")
    vitamin_c = models.FloatField(null=True, blank=True, verbose_name="Vitamin C (mg)")
    vitamin_d = models.FloatField(null=True, blank=True, verbose_name="Vitamin D (µg)")
    vitamin_e = models.FloatField(null=True, blank=True, verbose_name="Vitamin E (mg)")
    vitamin_b1 = models.FloatField(null=True, blank=True, verbose_name="Vitamin B1 (mg)")
    vitamin_b2 = models.FloatField(null=True, blank=True, verbose_name="Vitamin B2 (mg)")
    vitamin_b6 = models.FloatField(null=True, blank=True, verbose_name="Vitamin B6 (mg)")
    vitamin_b12 = models.FloatField(null=True, blank=True, verbose_name="Vitamin B12 (µg)")
    calcium = models.FloatField(null=True, blank=True, verbose_name="Kaltsiy (mg)")
    iron = models.FloatField(null=True, blank=True, verbose_name="Temir (mg)")
    magnesium = models.FloatField(null=True, blank=True, verbose_name="Magniy (mg)")
    zinc = models.FloatField(null=True, blank=True, verbose_name="Sink (mg)")
    potassium = models.FloatField(null=True, blank=True, verbose_name="Kaliy (mg)")
    
    # Sog'liq ko'rsatkichlari
    nutriscore_grade = models.CharField(max_length=1, blank=True, verbose_name="Nutri-Score (A-E)")
    nova_group = models.IntegerField(null=True, blank=True, verbose_name="NOVA guruhi (1-4)")
    ecoscore_grade = models.CharField(max_length=1, blank=True, verbose_name="Eco-Score (A-E)")
    
    # Allergenlar va ogohlantirishlar
    allergens = models.TextField(blank=True, verbose_name="Allergenlar")
    additives = models.TextField(blank=True, verbose_name="Qo'shimchalar (E raqamlar)")
    warnings = models.TextField(blank=True, verbose_name="Ogohlantirishlar")
    
    is_active = models.BooleanField(default=True, verbose_name="Faol")
    scan_count = models.PositiveIntegerField(default=0, verbose_name="Skanerlangan soni")
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = "Mahsulot"
        verbose_name_plural = "Mahsulotlar"
        ordering = ['-scan_count', 'name']
    
    def __str__(self):
        return f"{self.name} ({self.barcode})"


class ProductScanHistory(models.Model):
    """Mahsulot skanerlash tarixi"""
    user = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True, verbose_name="Foydalanuvchi")
    product = models.ForeignKey(Product, on_delete=models.CASCADE, related_name='scans', verbose_name="Mahsulot")
    scanned_at = models.DateTimeField(auto_now_add=True, verbose_name="Skanerlangan vaqt")
    
    class Meta:
        verbose_name = "Skanerlash tarixi"
        verbose_name_plural = "Skanerlash tarixi"
        ordering = ['-scanned_at']
    
    def __str__(self):
        return f"{self.product.name} - {self.scanned_at}"
