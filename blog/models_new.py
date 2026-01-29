from django.db import models
from django.contrib.auth.models import User


class BookPage(models.Model):
    book = models.ForeignKey('Book', on_delete=models.CASCADE, related_name='pages', verbose_name="Kitob")
    page_number = models.PositiveIntegerField(verbose_name="Sahifa raqami")
    text = models.TextField(verbose_name="Sahifa matni")

    class Meta:
        unique_together = ('book', 'page_number')
        ordering = ['page_number']
        verbose_name = "Kitob sahifasi"
        verbose_name_plural = "Kitob sahifalari"

    def __str__(self):
        return f"{self.book.title} - {self.page_number}-sahifa"

class Image(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    title = models.CharField(max_length=100, blank=True)
    image = models.ImageField(upload_to='images/')
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.title or self.image.name

class File(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    file = models.FileField(upload_to='files/')
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.file.name

class Feedback(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    name = models.CharField(max_length=100, blank=True)
    message = models.TextField()
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"{self.name or 'Anonim'}: {self.message[:50]}"


class Author(models.Model):
    """Adib modeli"""
    name = models.CharField(max_length=200, verbose_name="Adib ismi")
    bio = models.TextField(blank=True, verbose_name="Tarjimai hol")
    birth_year = models.IntegerField(null=True, blank=True, verbose_name="Tug'ilgan yili")
    death_year = models.IntegerField(null=True, blank=True, verbose_name="Vafot etgan yili")
    image = models.ImageField(upload_to='authors/', blank=True, null=True, verbose_name="Rasm")
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = "Adib"
        verbose_name_plural = "Adiblar"
        ordering = ['name']

    def __str__(self):
        return self.name



class Book(models.Model):
    """Kitob modeli"""
    author = models.ForeignKey(Author, on_delete=models.CASCADE, related_name='books', verbose_name="Adib")
    title = models.CharField(max_length=300, verbose_name="Kitob nomi")
    description = models.TextField(blank=True, verbose_name="Tavsif")
    year_written = models.IntegerField(null=True, blank=True, verbose_name="Yozilgan yili")
    file = models.FileField(upload_to='books/', blank=True, null=True, verbose_name="Fayl (PDF/DOCX)")
    content = models.TextField(blank=True, verbose_name="Kitob matni", help_text="Fayldan avtomatik o'qilgan to'liq matn")
    created_at = models.DateTimeField(auto_now_add=True)

    def save_pages_from_file(self):
        """Fayldan matnni sahifalarga bo'lib BookPage obyektlari sifatida saqlash"""
        from .models import BookPage
        if not self.file:
            return
        file_path = self.file.path
        file_name = self.file.name.lower()
        # Avval eski sahifalarni o'chir
        self.pages.all().delete()
        
        if file_name.endswith('.pdf'):
            # PDF: Har bir sahifa alohida saqlanadi (haqiqiy sahifa raqamlari)
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                BookPage.objects.create(book=self, page_number=i+1, text=text)
        
        elif file_name.endswith('.docx'):
            # DOCX: Page break'lar bo'yicha sahifalarga bo'lish
            from docx import Document
            from docx.oxml.ns import qn
            doc = Document(file_path)
            
            current_page_text = []
            page_number = 1
            
            for paragraph in doc.paragraphs:
                # Page break borligini tekshirish
                has_page_break = False
                for run in paragraph.runs:
                    if run._element.xml.find('w:br') != -1 and 'w:type="page"' in run._element.xml:
                        has_page_break = True
                        break
                    # lastRenderedPageBreak ham tekshirish
                    for child in run._element:
                        if child.tag == qn('w:lastRenderedPageBreak'):
                            has_page_break = True
                            break
                
                if has_page_break and current_page_text:
                    # Oldingi sahifani saqlash
                    BookPage.objects.create(
                        book=self, 
                        page_number=page_number, 
                        text='\n'.join(current_page_text)
                    )
                    page_number += 1
                    current_page_text = []
                
                current_page_text.append(paragraph.text)
            
            # Oxirgi sahifani saqlash
            if current_page_text:
                BookPage.objects.create(
                    book=self, 
                    page_number=page_number, 
                    text='\n'.join(current_page_text)
                )
        
        elif file_name.endswith('.txt'):
            # TXT: Butun matnni bitta sahifa sifatida saqlash (sahifa tushunchasi yo'q)
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            BookPage.objects.create(book=self, page_number=1, text=text)

    class Meta:
        verbose_name = "Kitob"
        verbose_name_plural = "Kitoblar"
        ordering = ['title']

    def __str__(self):
        return f"{self.title} - {self.author.name}"

    def extract_text_from_file(self):
        """Fayldan matnni o'qib olish"""
        if not self.file:
            return ""
        
        file_path = self.file.path
        file_name = self.file.name.lower()
        text = ""
        
        try:
            if file_name.endswith('.pdf'):
                text = self._extract_from_pdf(file_path)
            elif file_name.endswith('.docx'):
                text = self._extract_from_docx(file_path)
            elif file_name.endswith('.txt'):
                text = self._extract_from_txt(file_path)
        except Exception as e:
            text = f"Faylni o'qishda xatolik: {str(e)}"
        
        return text
    
    def _extract_from_pdf(self, file_path):
        """PDF fayldan matn o'qish"""
        from pypdf import PdfReader
        text = ""
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    
    def _extract_from_docx(self, file_path):
        """DOCX fayldan matn o'qish"""
        from docx import Document
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_txt(self, file_path):
        """TXT fayldan matn o'qish"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
